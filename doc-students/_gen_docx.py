"""สร้าง student-guide.docx จาก student-guide.md"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

OUT = os.path.join(os.path.dirname(__file__), "student-guide.docx")

# ── Palette ──────────────────────────────────────────────────────────
INDIGO  = RGBColor(0x4f, 0x46, 0xe5)
PURPLE  = RGBColor(0x7c, 0x3a, 0xed)
GREEN   = RGBColor(0x05, 0x96, 0x69)
MUTED   = RGBColor(0x64, 0x74, 0x8b)
TEXT    = RGBColor(0x0f, 0x17, 0x2a)
WHITE   = RGBColor(0xff, 0xff, 0xff)
LIGHT   = RGBColor(0xf1, 0xf5, 0xf9)
WARN_BG = RGBColor(0xff, 0xf7, 0xed)

FONT_TH = "Sarabun"
FONT_EN = "Sarabun"

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────
sec = doc.sections[0]
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.0)

# ── Default paragraph style ──────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = FONT_TH
style.font.size = Pt(11)
style.font.color.rgb = TEXT
style.paragraph_format.space_after  = Pt(4)
style.paragraph_format.space_before = Pt(0)


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_para(text="", bold=False, italic=False, size=11,
             color=None, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=4, font=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.name  = font or FONT_TH
        run.font.bold  = bold
        run.font.italic = italic
        run.font.size  = Pt(size)
        run.font.color.rgb = color or TEXT
    return p


def add_heading(text, level=1):
    sizes  = {1: 20, 2: 16, 3: 13}
    colors = {1: INDIGO, 2: INDIGO, 3: TEXT}
    p = add_para(text, bold=True, size=sizes.get(level, 12),
                 color=colors.get(level, TEXT),
                 space_before=12 if level == 1 else 8,
                 space_after=6)
    return p


def add_table(headers, rows, col_widths=None):
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "4f46e5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name  = FONT_TH
        run.font.bold  = True
        run.font.size  = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = tbl.rows[r_idx + 1]
        bg = "f8fafc" if r_idx % 2 == 0 else "ffffff"
        for c_idx, val in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = FONT_TH
            run.font.size = Pt(10)
            run.font.color.rgb = TEXT

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()


def add_note(text, emoji="📌", bg="fffbeb", border="fbbf24"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    p.paragraph_format.left_indent  = Cm(0.5)
    run = p.add_run(f"{emoji}  {text}")
    run.font.name   = FONT_TH
    run.font.size   = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0e)


def add_step(number, title, body_lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"  ขั้นตอนที่ {number} — ")
    r1.font.name  = FONT_TH
    r1.font.bold  = True
    r1.font.size  = Pt(11)
    r1.font.color.rgb = INDIGO
    r2 = p.add_run(title)
    r2.font.name  = FONT_TH
    r2.font.bold  = True
    r2.font.size  = Pt(11)
    r2.font.color.rgb = TEXT
    for line in body_lines:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent   = Cm(1.2)
        bp.paragraph_format.space_after   = Pt(2)
        bp.paragraph_format.space_before  = Pt(0)
        run = bp.add_run(line)
        run.font.name = FONT_TH
        run.font.size = Pt(10.5)
        run.font.color.rgb = TEXT


def add_qa(q, a):
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(6)
    pq.paragraph_format.space_after  = Pt(1)
    rq = pq.add_run(f"Q: {q}")
    rq.font.name  = FONT_TH
    rq.font.bold  = True
    rq.font.size  = Pt(10.5)
    rq.font.color.rgb = INDIGO

    pa = doc.add_paragraph()
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after  = Pt(6)
    pa.paragraph_format.left_indent  = Cm(0.5)
    ra = pa.add_run(f"A: {a}")
    ra.font.name = FONT_TH
    ra.font.size = Pt(10.5)
    ra.font.color.rgb = TEXT


# ════════════════════════════════════════════════════════════════════
#  COVER
# ════════════════════════════════════════════════════════════════════
add_para("ระบบประเมินสุขภาพจิตนักเรียน จังหวัดเลย",
         bold=True, size=13, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=20, space_after=4)
add_para("LEMCS", bold=True, size=34, color=INDIGO,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6)
add_para("Loei Educational MindCare System",
         italic=True, size=12, color=MUTED,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=20)
add_para("คู่มือการใช้งานสำหรับนักเรียน",
         bold=True, size=22, color=TEXT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=8)
add_para("", space_after=4)

doc.add_table(rows=0, cols=0)   # placeholder — skip; draw divider via para
p_div = doc.add_paragraph()
p_div.paragraph_format.space_before = Pt(0)
p_div.paragraph_format.space_after  = Pt(24)
run_div = p_div.add_run("─" * 60)
run_div.font.color.rgb = RGBColor(0xc7, 0xd2, 0xfe)
run_div.font.size = Pt(9)
p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  1. ระบบ LEMCS คืออะไร
# ════════════════════════════════════════════════════════════════════
add_heading("1. ระบบ LEMCS คืออะไร", 1)
add_para("LEMCS คือระบบคัดกรองสุขภาพจิตนักเรียนของจังหวัดเลย พัฒนาขึ้นเพื่อช่วยให้นักเรียนและครูสามารถตรวจสอบภาวะสุขภาพจิตได้อย่างสะดวก รวดเร็ว และเป็นส่วนตัว",
         space_after=8)

add_heading("แบบประเมินที่มีในระบบ", 2)
add_table(
    ["แบบประเมิน", "วัดอะไร", "กลุ่มอายุ", "จำนวนข้อ"],
    [
        ["🧠  ST-5",  "ความเครียด",               "อายุ 15 ปีขึ้นไป", "5 ข้อ"],
        ["🌱  CDI",   "ภาวะซึมเศร้าในเด็ก",       "อายุ 7–17 ปี",     "27 ข้อ"],
        ["💙  PHQ-A", "ภาวะซึมเศร้าในวัยรุ่น",    "อายุ 11–20 ปี",    "9 ข้อ"],
    ],
    col_widths=[3.5, 5, 4, 3],
)
add_note("ระบบแสดงเฉพาะแบบประเมินที่เหมาะกับอายุของนักเรียนแต่ละคนโดยอัตโนมัติ ไม่ต้องเลือกเอง", "ℹ️")

# ════════════════════════════════════════════════════════════════════
#  2. ก่อนเริ่มต้นใช้งาน
# ════════════════════════════════════════════════════════════════════
add_heading("2. ก่อนเริ่มต้นใช้งาน", 1)

add_heading("สิ่งที่ต้องเตรียม", 2)
for item in [
    "เลขบัตรประจำตัวประชาชน 13 หลัก",
    "วันเดือนปีเกิด (ปี พ.ศ.)",
    "รหัสประจำตัวนักเรียน (รับจากครูหรือโรงเรียน)",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(item)
    r.font.name = FONT_TH
    r.font.size = Pt(11)
    r.font.color.rgb = TEXT

doc.add_paragraph()
add_heading("อุปกรณ์และเบราว์เซอร์ที่รองรับ", 2)
add_table(
    ["อุปกรณ์", "เบราว์เซอร์ที่แนะนำ"],
    [
        ["สมาร์ทโฟน (Android / iOS)", "Google Chrome, Safari"],
        ["แท็บเล็ต",                   "Google Chrome, Safari"],
        ["คอมพิวเตอร์",                "Chrome, Edge, Firefox"],
    ],
    col_widths=[7, 8],
)

# ════════════════════════════════════════════════════════════════════
#  3. วิธีเข้าสู่ระบบ
# ════════════════════════════════════════════════════════════════════
add_heading("3. วิธีเข้าสู่ระบบ", 1)
add_para("เปิดเบราว์เซอร์แล้วไปที่ที่อยู่เว็บไซต์ที่ครูแจ้งให้ จากนั้นทำตามขั้นตอนดังนี้", space_after=8)

add_step(1, "กรอกเลขบัตรประจำตัวประชาชน", [
    "พิมพ์เลข 13 หลัก ลงในช่อง 'เลขบัตรประจำตัวประชาชน'",
    "ระบบจะจัดรูปแบบให้อัตโนมัติ เช่น 1-2345-67890-12-3",
    "แถบสีใต้ช่องจะเปลี่ยนเป็นสีเขียวเมื่อกรอกครบ 13 หลัก",
])

add_step(2, "เลือกวันเดือนปีเกิด (พ.ศ.)", [
    "เลือก วัน จากช่องแรก",
    "เลือก เดือน จากช่องกลาง",
    "เลือก ปี พ.ศ. จากช่องสุดท้าย",
])
add_note("ตัวอย่าง: เกิดวันที่ 15 มกราคม พ.ศ. 2550 → เลือก 15 / มกราคม / 2550", "📅")

add_step(3, "กรอกรหัสประจำตัวนักเรียน", [
    "พิมพ์รหัสนักเรียนที่ได้รับจากโรงเรียน",
    "รหัสเป็นตัวเลขเท่านั้น",
])

add_para("เมื่อกรอกครบ 3 ช่อง ปุ่ม 'เข้าสู่ระบบ' จะเปิดใช้งาน กดปุ่มแล้วรอสักครู่ ระบบจะพาไปหน้าหลักโดยอัตโนมัติ",
         space_before=6, space_after=4)

add_note("หากเข้าสู่ระบบไม่ได้: ตรวจสอบเลขบัตรประชาชน วันเดือนปีเกิด และรหัสนักเรียน หากยังไม่ได้ ติดต่อครูเพื่อขอตรวจสอบข้อมูล", "⚠️")

# ════════════════════════════════════════════════════════════════════
#  4. หน้าหลัก
# ════════════════════════════════════════════════════════════════════
add_heading("4. หน้าหลัก (Dashboard)", 1)

add_table(
    ["ส่วนประกอบ", "รายละเอียด"],
    [
        ["แถบสถานะการสำรวจ (สีเขียว)", "รอบการสำรวจเปิดอยู่ สามารถทำแบบประเมินได้ทันที"],
        ["แถบสถานะการสำรวจ (สีเหลือง)", "ยังไม่เปิดรอบ ต้องรอให้ครูเปิดก่อน ปุ่มจะเป็นสีเทา"],
        ["แบบประเมินที่เปิดให้ทำ", "รายการที่เหมาะกับอายุ กดปุ่มสีม่วง 'เริ่มเลย' เพื่อเริ่ม"],
        ["สุขภาพจิตของคุณ", "กราฟสรุปผลการประเมินย้อนหลัง"],
        ["ประวัติการเข้าร่วม", "รายการแบบประเมินที่เคยทำ พร้อมผลและวันที่"],
    ],
    col_widths=[5.5, 10],
)

# ════════════════════════════════════════════════════════════════════
#  5. การทำแบบประเมิน
# ════════════════════════════════════════════════════════════════════
add_heading("5. การทำแบบประเมิน", 1)

add_heading("เริ่มทำแบบประเมิน", 2)
add_para("กดปุ่ม 'เริ่มเลย' ที่แบบประเมินที่ต้องการ จากนั้นอ่านคำถามและเลือกคำตอบที่ตรงกับความรู้สึก", space_after=8)

add_heading("ระหว่างทำแบบประเมิน", 2)
add_table(
    ["องค์ประกอบ", "หน้าที่"],
    [
        ["แถบความคืบหน้า", "แสดงว่าทำไปกี่ข้อแล้วจากทั้งหมด"],
        ["วงกลมนำทาง (สีม่วงเข้ม)", "ข้อที่กำลังทำอยู่"],
        ["วงกลมนำทาง (สีม่วงอ่อน)", "ตอบแล้ว — กดเพื่อกลับไปแก้ไขได้"],
        ["วงกลมนำทาง (สีเทา)", "ยังไม่ได้ตอบ"],
        ["← ย้อนกลับ", "กลับไปแก้คำตอบข้อก่อนหน้า"],
        ["ข้ามไปข้อถัดไป →", "ข้ามไปข้อต่อไป (ต้องตอบข้อปัจจุบันก่อน)"],
    ],
    col_widths=[5.5, 10],
)

add_heading("คำแนะนำในการตอบ", 2)
add_table(
    ["แบบประเมิน", "ช่วงเวลาที่ใช้ประเมิน", "หมายเหตุ"],
    [
        ["ST-5 (ความเครียด)", "1 เดือนที่ผ่านมา", "เลือกระดับที่ตรงกับความรู้สึกมากที่สุด"],
        ["CDI (ซึมเศร้าในเด็ก)", "2 สัปดาห์ที่ผ่านมา", "เลือกประโยคที่ตรงกับความรู้สึกมากที่สุด"],
        ["PHQ-A (ซึมเศร้าวัยรุ่น)", "2 สัปดาห์ที่ผ่านมา", "เลือกความถี่ที่ตรงกับตัวเองมากที่สุด"],
    ],
    col_widths=[4.5, 4.5, 6.5],
)
add_note("ไม่มีคำตอบถูกหรือผิด ตอบตามความรู้สึกจริงของตัวเองอย่างตรงไปตรงมา", "💡")

add_heading("การส่งคำตอบ", 2)
for item in [
    "เมื่อตอบครบทุกข้อ ปุ่ม 'ส่งคำตอบและประเมินผล' จะสว่างขึ้น",
    "กดปุ่ม แล้วกด 'ส่งแบบประเมิน' ในกรอบยืนยัน",
    "รอสักครู่ ระบบจะแสดงผลทันที",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(item)
    r.font.name = FONT_TH; r.font.size = Pt(11); r.font.color.rgb = TEXT

add_note("หลังส่งแล้วไม่สามารถแก้ไขคำตอบได้ กรุณาอ่านคำถามให้ครบก่อนส่ง", "⚠️")

# ════════════════════════════════════════════════════════════════════
#  6. การดูผลการประเมิน
# ════════════════════════════════════════════════════════════════════
add_heading("6. การดูผลการประเมิน", 1)
add_para("หลังส่งคำตอบ ระบบจะแสดงหน้าผลการประเมินทันที ประกอบด้วย คะแนนที่ได้ ระดับ และคำแนะนำ หากผลอยู่ในระดับที่ต้องดูแลเป็นพิเศษ ระบบจะแสดงข้อมูลช่องทางขอความช่วยเหลือโดยอัตโนมัติ", space_after=8)
add_para("กด 'กลับหน้าหลัก' เพื่อทำแบบประเมินอื่นต่อ หรือดูประวัติการทำแบบประเมินก่อนหน้า")

# ════════════════════════════════════════════════════════════════════
#  7. ความหมายของผลการประเมิน
# ════════════════════════════════════════════════════════════════════
add_heading("7. ความหมายของผลการประเมิน", 1)

add_heading("ST-5 (ความเครียด) — คะแนน 0–15", 2)
add_table(
    ["ระดับ", "คะแนน", "ความหมาย"],
    [
        ["ปกติ",      "0–4",   "ความเครียดอยู่ในระดับที่จัดการได้"],
        ["เล็กน้อย",  "5–7",   "มีความเครียดบ้าง ควรหาวิธีผ่อนคลาย"],
        ["ปานกลาง",  "8–9",   "ความเครียดค่อนข้างสูง ควรปรึกษาครูแนะแนว"],
        ["สูง",       "10–15", "ความเครียดสูงมาก ควรพบผู้เชี่ยวชาญ"],
    ],
    col_widths=[3.5, 3, 9],
)

add_heading("CDI (ภาวะซึมเศร้าในเด็ก) — คะแนน 0–54", 2)
add_table(
    ["ระดับ", "คะแนน", "ความหมาย"],
    [
        ["ปกติ",       "0–14",      "ไม่มีภาวะซึมเศร้า"],
        ["ต้องดูแล",   "15 ขึ้นไป", "มีแนวโน้มภาวะซึมเศร้า ควรพบครูหรือผู้เชี่ยวชาญ"],
    ],
    col_widths=[3.5, 3.5, 8.5],
)

add_heading("PHQ-A (ภาวะซึมเศร้าในวัยรุ่น) — คะแนน 0–27", 2)
add_table(
    ["ระดับ", "คะแนน", "ความหมาย"],
    [
        ["ไม่มีอาการ",  "0–4",   "สุขภาพจิตดี"],
        ["เล็กน้อย",    "5–9",   "มีอาการเล็กน้อย ควรดูแลตัวเอง"],
        ["ปานกลาง",    "10–14", "ควรพูดคุยกับครูหรือผู้ปกครอง"],
        ["สูง",         "15–19", "ควรพบผู้เชี่ยวชาญด้านสุขภาพจิต"],
        ["รุนแรงมาก",  "20–27", "ควรพบแพทย์โดยเร็ว"],
    ],
    col_widths=[3.5, 3, 9],
)
add_note("ผลการประเมินใช้เพื่อคัดกรองเบื้องต้นเท่านั้น ไม่ใช่การวินิจฉัยโรค หากมีข้อสงสัยควรปรึกษาผู้เชี่ยวชาญ", "⚠️")

# ════════════════════════════════════════════════════════════════════
#  8. คำถามที่พบบ่อย
# ════════════════════════════════════════════════════════════════════
add_heading("8. คำถามที่พบบ่อย", 1)

add_qa("ทำไมแบบประเมินบางตัวไม่แสดงขึ้นมา?",
       "ระบบแสดงเฉพาะแบบประเมินที่เหมาะกับอายุ เช่น ST-5 สำหรับอายุ 15 ปีขึ้นไป หากอายุไม่ตรงเงื่อนไข แบบประเมินนั้นจะไม่แสดง")
add_qa("ทำไมปุ่มกดไม่ได้ ขึ้นว่า 'ยังไม่เปิด'?",
       "รอบการสำรวจยังไม่ถูกเปิดโดยครูหรือผู้ดูแลระบบ ให้รอการแจ้งจากครู")
add_qa("กรอกข้อมูลเข้าสู่ระบบแล้วขึ้นว่าไม่ถูกต้อง ทำอย่างไร?",
       "ตรวจสอบ (1) เลขบัตรประชาชน 13 หลักถูกต้อง (2) วันเดือนปีเกิดตรงกับที่แจ้งโรงเรียนไว้ (3) รหัสนักเรียนถูกต้อง หากยังไม่ได้ ติดต่อครูเพื่อตรวจสอบข้อมูลในระบบ")
add_qa("ทำแบบประเมินค้างไว้แล้วปิดหน้าต่าง ต้องเริ่มใหม่ไหม?",
       "ระบบบันทึกคำตอบชั่วคราวทุก 30 วินาที อย่างไรก็ตาม แนะนำให้ทำให้เสร็จในครั้งเดียวเพื่อความแม่นยำของผล")
add_qa("ผลการประเมินจะถูกเก็บเป็นความลับไหม?",
       "ข้อมูลทั้งหมดถูกเข้ารหัสและปกป้องตามมาตรฐาน PDPA ผลการประเมินจะถูกส่งให้ครูแนะแนวหรือผู้ดูแลเพื่อให้ความช่วยเหลือเท่านั้น")
add_qa("ทำแบบประเมินได้กี่ครั้ง?",
       "1 ครั้งต่อภาคเรียน ต่อประเภทแบบประเมิน ตามรอบที่โรงเรียนเปิด")
add_qa("ใช้มือถือทำได้ไหม?",
       "ได้ ระบบออกแบบสำหรับมือถือ (Mobile First) รองรับทุกขนาดหน้าจอตั้งแต่ 375px ขึ้นไป")

# ════════════════════════════════════════════════════════════════════
#  9. ขอความช่วยเหลือ
# ════════════════════════════════════════════════════════════════════
add_heading("9. ขอความช่วยเหลือ", 1)
add_para("หากรู้สึกไม่ดีหรือต้องการพูดคุย สามารถติดต่อได้ที่", space_after=6)

add_table(
    ["ช่องทาง", "รายละเอียด"],
    [
        ["🆘  สายด่วนสุขภาพจิต",           "1323 (ตลอด 24 ชั่วโมง)"],
        ["👩‍🏫  ครูแนะแนว",                "ติดต่อที่โรงเรียนของนักเรียน"],
        ["🏥  สำนักงานสาธารณสุขจังหวัดเลย", "โทร 042-811-522"],
    ],
    col_widths=[7, 8.5],
)

# ── Footer ───────────────────────────────────────────────────────────
doc.add_paragraph()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(16)
rf = p_footer.add_run("จัดทำโดยทีมพัฒนา LEMCS — สำนักงานศึกษาธิการจังหวัดเลย  |  ปรับปรุงล่าสุด: เมษายน 2569")
rf.font.name = FONT_TH
rf.font.size = Pt(9)
rf.font.italic = True
rf.font.color.rgb = MUTED

doc.save(OUT)
print(f"OK: {OUT}")
