# -*- coding: utf-8 -*-
"""Generate superadmin-manual.docx from structured content."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "superadmin-manual.docx")

# ─── Color Palette ───────────────────────────────────────────────
C_NAVY    = RGBColor(0x0b, 0x15, 0x26)
C_TEAL    = RGBColor(0x00, 0x99, 0x8a)
C_BLUE    = RGBColor(0x1e, 0x6f, 0xbf)
C_RED     = RGBColor(0xc0, 0x39, 0x2b)
C_GREEN   = RGBColor(0x1e, 0x8a, 0x44)
C_ORANGE  = RGBColor(0xe6, 0x7e, 0x22)
C_PURPLE  = RGBColor(0x7d, 0x3c, 0x98)
C_YELLOW  = RGBColor(0xb7, 0x7a, 0x00)
C_DGRAY   = RGBColor(0x2c, 0x3e, 0x50)
C_LGRAY   = RGBColor(0x85, 0x92, 0x9e)
C_BLACK   = RGBColor(0x1a, 0x1a, 0x2e)
C_WHITE   = RGBColor(0xff, 0xff, 0xff)

# ─── Helper: set paragraph shading ──────────────────────────────
def shade_paragraph(paragraph, hex_color):
    """Add background shading to paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_cell_bg(cell, hex_color):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), kwargs.get(edge, 'single'))
        tag.set(qn('w:sz'), str(kwargs.get('sz', 4)))
        tag.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ─── Document Setup ──────────────────────────────────────────────
doc = Document()

# Page margins
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin   = Cm(2.5)
sec.right_margin  = Cm(2.5)
sec.top_margin    = Cm(2.0)
sec.bottom_margin = Cm(2.0)

# Core styles
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'TH Sarabun New'
normal.font.size = Pt(14)
normal.paragraph_format.space_after = Pt(4)

h1s = styles['Heading 1']
h1s.font.name = 'TH Sarabun New'
h1s.font.size = Pt(22)
h1s.font.bold = True
h1s.font.color.rgb = C_NAVY

h2s = styles['Heading 2']
h2s.font.name = 'TH Sarabun New'
h2s.font.size = Pt(18)
h2s.font.bold = True
h2s.font.color.rgb = C_BLUE

h3s = styles['Heading 3']
h3s.font.name = 'TH Sarabun New'
h3s.font.size = Pt(15)
h3s.font.bold = True
h3s.font.color.rgb = C_TEAL

# ─── Helper Functions ────────────────────────────────────────────
def add_heading(doc, text, level=1, color=None):
    style = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3'}.get(level, 'Heading 2')
    p = doc.add_heading(text, level=level)
    p.style = doc.styles[style]
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(doc, text, bold=False, italic=False, color=None, size=14, indent=0, space_before=0, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def add_bullet(doc, text, level=0, color=None, size=13):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_numbered(doc, text, color=None, size=13):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_screenshot_box(doc, caption):
    """Add a dashed screenshot placeholder box."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0)
    # Camera icon + caption
    run = p.add_run(f"  📸  {caption}")
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = C_LGRAY
    shade_paragraph(p, 'F0F4F8')
    p.paragraph_format.left_indent = Cm(0.3)
    # blank line below for image space
    blank = doc.add_paragraph()
    blank.paragraph_format.space_before = Pt(0)
    blank.paragraph_format.space_after  = Pt(0)
    blank_run = blank.add_run(" " * 5)
    blank_run.font.size = Pt(60)   # tall blank space
    shade_paragraph(blank, 'E8EFF7')
    # bottom label
    bot = doc.add_paragraph()
    bot_run = bot.add_run("  [ วางภาพแคปหน้าจอที่นี่ ]")
    bot_run.font.name = 'TH Sarabun New'
    bot_run.font.size = Pt(10)
    bot_run.font.italic = True
    bot_run.font.color.rgb = C_LGRAY
    shade_paragraph(bot, 'E8EFF7')
    bot.paragraph_format.space_after = Pt(8)

def add_note_box(doc, text, bg='FFF9E6', color=None):
    """Add a colored note/warning box."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'TH Sarabun New'
    run.font.size = Pt(12.5)
    if color:
        run.font.color.rgb = color
    shade_paragraph(p, bg)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)

def add_table(doc, headers, rows, col_widths=None, header_bg='1B4F72', header_fg=C_WHITE):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.name = 'TH Sarabun New'
        cell.paragraphs[0].runs[0].font.size = Pt(13)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = header_fg
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'F2F7FB'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.name = 'TH Sarabun New'
            cell.paragraphs[0].runs[0].font.size = Pt(12.5)
            set_cell_bg(cell, bg)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('─' * 80)
    run.font.color.rgb = RGBColor(0xCC, 0xD9, 0xE8)
    run.font.size = Pt(8)

def add_page_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════
cover_p = doc.add_paragraph()
cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_p.paragraph_format.space_before = Pt(40)
shade_paragraph(cover_p, '0B1526')
r = cover_p.add_run("LEMCS")
r.font.name = 'TH Sarabun New'
r.font.size = Pt(42)
r.font.bold = True
r.font.color.rgb = RGBColor(0x00, 0xC9, 0xB1)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub_p, '0B1526')
r2 = sub_p.add_run("Loei Educational MindCare System")
r2.font.name = 'TH Sarabun New'
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x8F, 0xA3, 0xBC)

for line in [
    ("คู่มือการใช้งานสำหรับ Superadmin", Pt(24), RGBColor(0xe8, 0xf0, 0xfe), True),
    ("ผู้ดูแลระบบระดับจังหวัด — สำนักงานศึกษาธิการจังหวัดเลย", Pt(15), RGBColor(0x8F, 0xA3, 0xBC), False),
    ("เวอร์ชัน 1.0  ·  ปีการศึกษา 2568", Pt(13), RGBColor(0x8F, 0xA3, 0xBC), False),
]:
    lp = doc.add_paragraph()
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_paragraph(lp, '0B1526')
    lr = lp.add_run(line[0])
    lr.font.name = 'TH Sarabun New'
    lr.font.size = line[1]
    lr.font.color.rgb = line[2]
    lr.font.bold = line[3]

spacer = doc.add_paragraph()
spacer.paragraph_format.space_after = Pt(40)
shade_paragraph(spacer, '0B1526')

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (static)
# ════════════════════════════════════════════════════════════════
add_heading(doc, "สารบัญ", level=1, color=C_NAVY)
toc_items = [
    ("บทที่ 1", "บทนำและภาพรวมระบบ"),
    ("บทที่ 2", "การเข้าสู่ระบบ"),
    ("บทที่ 3", "หน้า Dashboard หลัก"),
    ("บทที่ 4", "การจัดการรอบสำรวจ (Survey Rounds)"),
    ("บทที่ 5", "ระบบแจ้งเตือนความเสี่ยง (Alerts)"),
    ("บทที่ 6", "รายงานรายบุคคล (Reports)"),
    ("บทที่ 7", "ข้อมูลนักเรียน (Students)"),
    ("บทที่ 8", "การจัดการผู้ใช้งาน (Users)"),
    ("บทที่ 9", "ข้อมูลโรงเรียน (Schools)"),
    ("บทที่ 10", "บันทึกการเข้าถึงข้อมูล (Audit Logs)"),
    ("บทที่ 11", "การตั้งค่าระบบ (Settings)"),
    ("บทที่ 12", "ขั้นตอนการทำงานทั่วไป"),
    ("ภาคผนวก", "คำถามที่พบบ่อย"),
]
for ch, title in toc_items:
    tp = doc.add_paragraph()
    r_ch = tp.add_run(f"{ch}  ")
    r_ch.font.name = 'TH Sarabun New'
    r_ch.font.size = Pt(13)
    r_ch.font.bold = True
    r_ch.font.color.rgb = C_BLUE
    r_ti = tp.add_run(title)
    r_ti.font.name = 'TH Sarabun New'
    r_ti.font.size = Pt(13)
    tp.paragraph_format.space_after = Pt(3)
    tp.paragraph_format.left_indent = Cm(0.3)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 1: Introduction
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 1: บทนำและภาพรวมระบบ", level=1)

add_heading(doc, "1.1 ระบบ LEMCS คืออะไร", level=2)
add_para(doc,
    "LEMCS (Loei Educational MindCare System) เป็นระบบประเมินสุขภาพจิตนักเรียนระดับดิจิทัล "
    "ครอบคลุมนักเรียนระดับอนุบาล–มัธยมปลาย ในจังหวัดเลยกว่า 100,000 คน "
    "รองรับ 4 สังกัดหลัก ได้แก่ สพฐ., สพม., สอศ. (อาชีวะ), สช. (เอกชน) และ สกร."
)

add_heading(doc, "แบบประเมินที่ใช้ในระบบ", level=3)
add_table(doc,
    headers=["แบบประเมิน", "ชื่อเต็ม", "กลุ่มอายุ", "คะแนนเต็ม"],
    rows=[
        ["ST-5", "แบบประเมินความเครียด", "อายุ ≥ 15 ปี", "0–15"],
        ["PHQ-A", "Patient Health Questionnaire-Adolescent", "อายุ 11–20 ปี", "0–27"],
        ["CDI", "Children's Depression Inventory", "อายุ 7–17 ปี", "0–54"],
    ],
    col_widths=[2.5, 6.5, 3.5, 3.0],
)

add_heading(doc, "1.2 บทบาท Superadmin", level=2)
add_para(doc,
    "Superadmin คือผู้ดูแลระบบระดับจังหวัด (สำนักงานศึกษาธิการจังหวัดเลย) มีสิทธิ์ดูข้อมูลทุกสังกัด "
    "ทุกอำเภอ และทุกโรงเรียนโดยไม่ถูกจำกัดขอบเขต"
)

add_heading(doc, "ตารางสิทธิ์การใช้งาน", level=3)
add_table(doc,
    headers=["หมวดงาน", "ดูข้อมูล", "แก้ไข/ดำเนินการ"],
    rows=[
        ["Dashboard & Analytics", "✅", "✅ (Export/Download)"],
        ["รอบสำรวจ", "✅", "✅ (เปิด/ปิด/ยกเลิก/ลบ)"],
        ["ระบบแจ้งเตือน (Alerts)", "✅", "✅ (อัปเดตสถานะ)"],
        ["รายงานรายบุคคล", "✅", "✅ (Export)"],
        ["ข้อมูลนักเรียน", "✅", "✅ (แก้ไขข้อมูลพื้นฐาน/NID)"],
        ["ผู้ใช้งานระบบ", "✅", "✅ (สร้าง/แก้ไข/Reset)"],
        ["ข้อมูลโรงเรียน", "✅", "❌ (เฉพาะ systemadmin)"],
        ["Audit Logs", "✅", "❌ (ดูอย่างเดียว)"],
        ["นำเข้าข้อมูล (Import)", "❌", "❌ (เฉพาะ systemadmin)"],
        ["โครงสร้างองค์กร", "❌", "❌ (เฉพาะ systemadmin)"],
    ],
    col_widths=[6.5, 3.0, 6.5],
)

add_heading(doc, "1.3 โครงสร้างลำดับชั้น", level=2)
for line in [
    "systemadmin    (ทีมพัฒนา — สิทธิ์สูงสุด)",
    "    └── superadmin    (ศึกษาธิการจังหวัด)  ◀ ท่านอยู่ระดับนี้",
    "            └── commissionadmin    (ศึกษาธิการอำเภอ / สังกัด)",
    "                    └── schooladmin    (ผู้บริหารโรงเรียน)",
    "                            └── student    (นักเรียน)",
]:
    p = doc.add_paragraph()
    r = p.add_run(line)
    r.font.name = 'Courier New'
    r.font.size = Pt(11)
    r.font.color.rgb = C_DGRAY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(1)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 2: Login
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 2: การเข้าสู่ระบบ", level=1)

add_heading(doc, "2.1 URL สำหรับ Admin", level=2)
add_table(doc,
    headers=["สภาพแวดล้อม", "URL"],
    rows=[
        ["Production", "https://lemcs.loeitech.ac.th/admin-login"],
        ["Dev/Staging", "https://dev.lemcs.loeitech.ac.th/admin-login"],
    ],
    col_widths=[4.0, 12.0],
)

add_heading(doc, "2.2 ขั้นตอนการเข้าสู่ระบบ", level=2)
for step in [
    "เปิดเบราว์เซอร์และไปที่ URL ด้านบน",
    "กรอก ชื่อผู้ใช้ (Username) และ รหัสผ่าน (Password)",
    "กดปุ่ม 'เข้าสู่ระบบ'",
    "ระบบจะพาไปยังหน้า /admin/dashboard โดยอัตโนมัติ",
]:
    add_numbered(doc, step)

add_note_box(doc, "⏱ หมายเหตุ: Session ของ Admin มีอายุ 8 ชั่วโมง หากไม่ได้ใช้งาน ระบบจะ Logout อัตโนมัติ")
add_screenshot_box(doc, "หน้า Admin Login — แสดงช่องกรอก Username/Password และปุ่มเข้าสู่ระบบ")

add_heading(doc, "2.3 การออกจากระบบ", level=2)
add_bullet(doc, "คลิกชื่อผู้ใช้ที่มุมขวาบน → เลือก 'ออกจากระบบ'")
add_bullet(doc, "หรือ Session หมดอายุอัตโนมัติหลัง 8 ชั่วโมง")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 3: Dashboard
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 3: หน้า Dashboard หลัก", level=1)
add_para(doc, "เส้นทาง: /admin/dashboard", color=C_LGRAY, size=12)
add_para(doc,
    "Dashboard คือหน้าหลักที่แสดงข้อมูลภาพรวมทั้งหมดแบบ Real-time แบ่งออกเป็น 4 Section "
    "พร้อม FilterBar ที่ควบคุมข้อมูลทุก Section พร้อมกัน"
)

add_screenshot_box(doc, "หน้า Dashboard ทั้งหน้า — แสดง Layout ของ 4 Section (A/B/C/D)")

add_heading(doc, "3.1 FilterBar — การกรองข้อมูล (9 พารามิเตอร์)", level=2)
add_para(doc,
    "FilterBar อยู่ด้านบนสุดของ Dashboard ใช้ควบคุมข้อมูลที่แสดงใน ทุก Chart และทุก Section พร้อมกัน"
)
add_table(doc,
    headers=["พารามิเตอร์", "คำอธิบาย"],
    rows=[
        ["รอบสำรวจ", "เลือกดูเฉพาะรอบที่ต้องการ"],
        ["สังกัด (Affiliation)", "สพฐ. / สพม. / สอศ. / สกร. / สช."],
        ["อำเภอ (District)", "อัปเดตอัตโนมัติตามสังกัดที่เลือก"],
        ["โรงเรียน (School)", "อัปเดตอัตโนมัติตามอำเภอที่เลือก"],
        ["ประเภทแบบประเมิน", "ST-5 / PHQ-A / CDI"],
        ["ชั้นเรียน (Grade)", "ป.1–ป.6, ม.1–ม.6, ปวช./ปวส."],
        ["เพศ (Gender)", "ชาย / หญิง"],
        ["วันที่เริ่ม", "วันที่เริ่มต้นที่ต้องการดู"],
        ["วันที่สิ้นสุด", "วันที่สิ้นสุดที่ต้องการดู"],
    ],
    col_widths=[5.0, 11.0],
)
add_note_box(doc, "💡 เคล็ดลับ: Dropdown สังกัด → อำเภอ → โรงเรียน เชื่อมกันเป็น Hierarchical cascade เลือกสังกัดแล้วอำเภอจะอัปเดตเองอัตโนมัติ")

add_screenshot_box(doc, "FilterBar — แสดง Dropdown ทั้ง 9 ตัวและปุ่ม Apply/Reset")

# Section A
add_heading(doc, "3.2 Section A — สรุปภาพรวมผู้บริหาร (Refresh ทุก 60 วินาที)", level=2)

add_heading(doc, "① StatsCards — ตัวเลขหลัก 4 ตัว", level=3)
for item in [
    "นักเรียนทั้งหมด — จำนวนนักเรียนในระบบ ณ ปัจจุบัน",
    "ทำแล้วในรอบนี้ — จำนวน + % coverage ของรอบที่เปิดอยู่",
    "ความเสี่ยง moderate+ — นักเรียนที่ต้องติดตามโดยด่วน",
    "Alert รอดำเนินการ — เคสที่ยังไม่มีผู้รับไปดูแล",
]:
    add_bullet(doc, item)

add_screenshot_box(doc, "StatsCards ทั้ง 4 การ์ด — ตัวเลขหลักที่เห็นได้ทันที")

add_heading(doc, "② RiskFunnelChart — Funnel 4 ขั้น", level=3)
add_para(doc,
    "แผนภูมิ Funnel แสดงจำนวนนักเรียนที่ลดหลั่นในแต่ละขั้นตอน:"
)
p = doc.add_paragraph()
r = p.add_run("ลงทะเบียน  →  ทำแบบประเมินแล้ว  →  พบความเสี่ยง  →  ดูแลแล้ว\n 100,000             45,000                    1,200              980")
r.font.name = 'Courier New'
r.font.size = Pt(11)
r.font.color.rgb = C_TEAL
p.paragraph_format.left_indent = Cm(1)
add_para(doc, "ใช้ระบุว่ากระบวนการดูแลมีช่องว่างตรงไหน เช่น นักเรียนทำแบบประเมินน้อยหรือเคสเสี่ยงยังไม่ได้รับการดูแล")

add_screenshot_box(doc, "RiskFunnelChart — แสดง 4 ขั้น พร้อมตัวเลขในแต่ละขั้น")

add_heading(doc, "③ InsightPanel — ข้อสรุปอัตโนมัติ", level=3)
for item in [
    "โรงเรียน Top 3 ที่มีอัตราความเสี่ยงสูงสุด",
    "อำเภอที่มีนักเรียนยังไม่ทำแบบประเมินมากที่สุด",
    "แบบประเมินที่พบความเสี่ยง moderate+ มากที่สุด",
]:
    add_bullet(doc, item)

add_heading(doc, "④ AffiliationStudentStats — Bar Chart แยกสังกัด", level=3)
add_bullet(doc, "แสดงจำนวนนักเรียนทั้งหมดแยก สพฐ. / สอศ. / สช. / สกร.")
add_bullet(doc, "เปรียบเทียบ % ที่ทำแบบประเมินแล้วในแต่ละสังกัด")

add_screenshot_box(doc, "Section A ทั้งหมด — InsightPanel และ AffiliationStudentStats")

# Section B
add_heading(doc, "3.3 Section B — ศูนย์ปฏิบัติการ (Refresh ทุก 30 วินาที)", level=2)

add_heading(doc, "① AlertStatusSummary", level=3)
add_table(doc,
    headers=["สถานะ", "สี", "ความหมาย"],
    rows=[
        ["รอดำเนินการ", "🔴 แดง", "ยังไม่มีผู้รับเคส"],
        ["กำลังดูแล",   "🟡 เหลือง", "ครู/แนะแนวรับเคสแล้ว"],
        ["เสร็จสิ้น",   "🟢 เขียว", "บันทึกผลการดูแลแล้ว"],
    ],
    col_widths=[4.0, 4.0, 8.0],
)

add_heading(doc, "② RecentAlerts — ตาราง 8 รายการล่าสุด", level=3)
for item in [
    "ชื่อนักเรียน (masked บางส่วน ตาม PDPA) · ชั้น/ห้อง · โรงเรียน",
    "ประเภท Alert: CDI Clinical (≥15) หรือ PHQ-A Suicide Risk (Q9≥1 / BQ=true)",
    "วันที่-เวลาที่ส่งแบบประเมิน · สถานะปัจจุบัน",
    "ปุ่ม: เปลี่ยนสถานะ → รอดำเนินการ / กำลังดูแล / เสร็จสิ้น",
    "ปุ่ม: ดูประวัตินักเรียนทั้งหมด (drill-down รายบุคคล)",
]:
    add_bullet(doc, item)

add_heading(doc, "③ Quick Actions — ปุ่มลัด 4 ปุ่ม", level=3)
add_table(doc,
    headers=["ปุ่ม", "ปลายทาง"],
    rows=[
        ["รายชื่อนักเรียน", "/admin/students"],
        ["Export รายงาน",   "/admin/reports"],
        ["นำเข้าข้อมูล",   "/admin/import"],
        ["จัดการรอบสำรวจ", "/admin/survey-rounds"],
    ],
    col_widths=[7.0, 9.0],
)

add_screenshot_box(doc, "Section B — AlertStatusSummary, RecentAlerts (8 รายการ), และ Quick Actions")

# Section C
add_heading(doc, "3.4 Section C — วิเคราะห์เชิงลึก (5 กราฟ)", level=2)

charts = [
    ("TrendChart — แนวโน้ม 6 เดือนล่าสุด",
     ["แกน X: เดือน (ต.ค.–มี.ค. หรือตามช่วงที่เลือก)",
      "แกน Y: จำนวนที่พบความเสี่ยง",
      "3 เส้นแยกสี: 🟡 ST-5 · 🔵 PHQ-A · 🟢 CDI",
      "ดูได้ทันทีว่าช่วงไหนมีความเสี่ยงพุ่งสูง (เช่น ช่วงสอบ)"]),
    ("MoMDeltaChart — เปรียบเทียบเดือนนี้ vs เดือนก่อน",
     ["Bar 🔴 แดง = ความเสี่ยงเพิ่มขึ้น (น่าเป็นห่วง)",
      "Bar 🟢 เขียว = ความเสี่ยงลดลง (ดีขึ้น)",
      "หน่วยเป็น % เช่น PHQ-A +12% หรือ CDI −5%"]),
    ("SeverityChart — Doughnut สัดส่วนระดับความเสี่ยง",
     ["🟢 ปกติ — ไม่ต้องดำเนินการเพิ่ม",
      "🟡 เล็กน้อย — ให้คำแนะนำ",
      "🟠 ปานกลาง — ครูแนะแนวดูแล",
      "🔴 รุนแรง — ต้องดำเนินการทันที",
      "ตัวเลขตรงกลาง = จำนวนนักเรียนที่ประเมินแล้วทั้งหมด"]),
    ("AssessmentTypeChart — Stacked Bar แยกประเภท",
     ["แกน X: ST-5 / PHQ-A / CDI",
      "แต่ละ Bar = สัดส่วน % ทุกระดับรวม 100%",
      "เห็นได้ทันทีว่าแบบประเมินไหนพบความเสี่ยงรุนแรงมากที่สุด"]),
    ("RiskProgressBars — แถบ Progress 4 ระดับ",
     ["แสดง N คน (xx%) ในแต่ละระดับ",
      "กดที่แถบเพื่อ drill-down ดูรายชื่อนักเรียนในระดับนั้น"]),
]

for chart_title, bullets in charts:
    add_heading(doc, chart_title, level=3)
    for b in bullets:
        add_bullet(doc, b)

add_screenshot_box(doc, "Section C — แสดงกราฟทั้ง 5 ตัว (Trend/MoM/Severity/Type/Progress)")

# Section D
add_heading(doc, "3.5 Section D — เปรียบเทียบองค์กร", level=2)
add_heading(doc, "OrgCompareChart — 100% Stacked Bar", level=3)
for item in [
    "ทุก Bar ยาว 100% เสมอ — เปรียบเทียบสัดส่วน % ได้ยุติธรรม ไม่ถูกบิดเบือนโดยขนาดโรงเรียน",
    "โรงเรียนเล็ก 300 คน และโรงเรียนใหญ่ 3,000 คน สามารถเปรียบเทียบกันได้",
    "สีแต่ละส่วน = ระดับความเสี่ยง (เขียว/เหลือง/ส้ม/แดง)",
    "มองเห็นทันทีว่าองค์กรใดมี % สีแดง+ส้มมากกว่ากัน → ต้องให้ทรัพยากรเพิ่ม",
]:
    add_bullet(doc, item)

add_heading(doc, "Auto-grouping — ระบบจัดกลุ่มอัตโนมัติ", level=3)
add_table(doc,
    headers=["FilterBar ที่เลือก", "ระบบจัดกลุ่มโดย", "ตัวอย่าง"],
    rows=[
        ["ไม่กรอง (ทั้งจังหวัด)", "สังกัด", "สพฐ. / สอศ. / สช. / สกร. (4 bars)"],
        ["เลือก Affiliation", "อำเภอ", "14 อำเภอในจังหวัดเลย"],
        ["เลือก District", "โรงเรียน", "ทุกโรงเรียนในอำเภอที่เลือก"],
    ],
    col_widths=[5.0, 4.5, 6.5],
)

add_screenshot_box(doc, "Section D — OrgCompareChart 100% Stacked Bar เปรียบเทียบองค์กร")

add_heading(doc, "3.6 การส่งออกจาก Dashboard", level=2)
add_table(doc,
    headers=["ปุ่ม", "ผลลัพธ์"],
    rows=[
        ["📄 PDF", "ดาวน์โหลด lemcs_report.pdf พร้อม Chart ทั้งหมด"],
        ["📊 Excel", "ดาวน์โหลด lemcs_report.xlsx ข้อมูลตาราง"],
        ["🖼️ ↓ PNG (ทุก Chart)", "ดาวน์โหลดกราฟแยกชิ้น"],
        ["🖨️ Print", "พิมพ์หน้า Dashboard (ซ่อน Sidebar อัตโนมัติ)"],
    ],
    col_widths=[5.5, 10.5],
)

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 4: Survey Rounds
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 4: การจัดการรอบสำรวจ", level=1)
add_para(doc, "เส้นทาง: /admin/survey-rounds", color=C_LGRAY, size=12)
add_para(doc,
    "รอบสำรวจ (Survey Round) คือช่วงเวลาที่นักเรียนสามารถเข้าทำแบบประเมินได้ "
    "ต้องเปิดรอบก่อนนักเรียนถึงจะทำแบบประเมินได้"
)

add_screenshot_box(doc, "หน้า Survey Rounds — แสดงตารางรอบสำรวจทั้งหมดพร้อมสถานะ")

add_heading(doc, "4.1 สถานะของรอบสำรวจ", level=2)
add_table(doc,
    headers=["สถานะ", "สัญลักษณ์", "ความหมาย"],
    rows=[
        ["เปิด (Open)", "🟢 กะพริบ", "นักเรียนทำแบบประเมินได้"],
        ["ปิด (Closed)", "⚫", "สิ้นสุดการเก็บข้อมูล รวมในรายงาน"],
        ["ยกเลิก (Cancelled)", "⏸", "หยุดชั่วคราว ข้อมูลไม่นับในรายงาน"],
    ],
    col_widths=[4.5, 4.0, 7.5],
)

add_heading(doc, "4.2 การเปิดรอบสำรวจใหม่", level=2)
for step in [
    "คลิกปุ่ม '➕ เปิดรอบใหม่'",
    "กรอก ชื่อรอบ (Label) เช่น 'ภาคเรียน 1/2568'",
    "เลือก ปีการศึกษา (Academic Year) และ ภาคเรียน (1 หรือ 2)",
    "คลิก 'เปิดรอบ' → ระบบสร้าง SurveyRound สถานะ 'open'",
    "นักเรียนจะเห็น Banner สีเขียวบน Dashboard ของตนเอง",
]:
    add_numbered(doc, step)

add_note_box(doc, "⚠️ ข้อควรระวัง: ไม่ควรเปิดหลายรอบพร้อมกัน เพราะจะทำให้นักเรียนสับสน", bg='FFF0F0')
add_screenshot_box(doc, "Modal เปิดรอบสำรวจใหม่ — แสดงช่องกรอก Label/ปีการศึกษา/ภาคเรียน")

add_heading(doc, "4.3 การปิดรอบสำรวจ", level=2)
for step in [
    "คลิกปุ่ม 'ปิดรอบ' ที่แถวของรอบที่ต้องการ",
    "ยืนยันการปิด",
]:
    add_numbered(doc, step)
add_para(doc, "ผลที่ตามมา:")
add_bullet(doc, "✅ ข้อมูลทั้งหมดยังคงอยู่ในระบบ")
add_bullet(doc, "✅ รวมในสถิติและรายงาน")
add_bullet(doc, "❌ นักเรียนไม่สามารถส่งแบบประเมินใหม่ได้")

add_heading(doc, "4.4 การยกเลิกรอบสำรวจ (Cancel)", level=2)
for step in [
    "คลิกปุ่ม '⏸ ยกเลิกรอบ'",
    "ยืนยันการยกเลิก",
]:
    add_numbered(doc, step)
add_para(doc, "ผลที่ตามมา:")
add_bullet(doc, "✅ ข้อมูลทั้งหมดยังคงอยู่ในฐานข้อมูล (Audit Trail)")
add_bullet(doc, "✅ Alert ที่สร้างไว้ยังคงอยู่ (เพื่อความปลอดภัยของนักเรียน)")
add_bullet(doc, "❌ ไม่นับในสถิติและรายงานทั่วไป")
add_bullet(doc, "❌ นักเรียนไม่สามารถส่งแบบประเมินใหม่ได้")
add_note_box(doc, "💡 ใช้ Cancel เมื่อ: เปิดรอบผิดพลาด หรือต้องการระงับชั่วคราวโดยยังเก็บข้อมูลไว้")

add_heading(doc, "4.5 การลบรอบสำรวจ (⚠️ ข้อมูลจะหายถาวร)", level=2)
for step in [
    "คลิกปุ่ม '🗑️ ลบรอบ'",
    "ระบบแสดง Modal ยืนยัน — ต้องพิมพ์ 'ลบถาวร' เพื่อยืนยัน",
    "ข้อมูลแบบประเมินทั้งหมดในรอบนั้นจะถูกลบถาวร",
]:
    add_numbered(doc, step)

add_note_box(doc,
    "🚨 คำเตือน: ใช้การลบเฉพาะกรณีที่รอบเปิดผิดพลาดและยังไม่มีนักเรียนทำแบบประเมิน "
    "หรือต้องการล้างข้อมูล Test เท่านั้น — ไม่สามารถกู้คืนได้",
    bg='FFE8E8'
)
add_screenshot_box(doc, "Modal ยืนยันการลบรอบ — แสดงช่องพิมพ์ 'ลบถาวร' และคำเตือน")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 5: Alerts
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 5: ระบบแจ้งเตือนความเสี่ยง", level=1)
add_para(doc, "เส้นทาง: /admin/alerts", color=C_LGRAY, size=12)

add_heading(doc, "5.1 เงื่อนไขการสร้าง Alert", level=2)
add_table(doc,
    headers=["แบบประเมิน", "เงื่อนไข", "ระดับ"],
    rows=[
        ["PHQ-A", "Q9 ≥ 1 คะแนน", "🚨 Critical (Suicide Risk)"],
        ["PHQ-A", "BQ1 หรือ BQ2 = true", "🚨 Critical (Suicide Risk)"],
        ["CDI", "คะแนนรวม ≥ 15", "🔴 Urgent (Clinical)"],
        ["ST-5/PHQ-A/CDI", "ความเสี่ยงระดับ moderate+", "⚠️ Warning"],
    ],
    col_widths=[3.5, 7.0, 5.5],
)
add_note_box(doc,
    "⚡ สำคัญ: Alert ประเภท Suicide Risk จะแจ้งเตือนใน Dashboard ทุกระดับพร้อมกัน "
    "(ครู → ผู้บริหารอำเภอ → จังหวัด) โดยทันที (Synchronous) — ไม่มีการหน่วงเวลา",
    bg='FFF0F0'
)

add_heading(doc, "5.2 การกรอง Alert", level=2)
add_table(doc,
    headers=["ตัวกรอง", "ตัวเลือก"],
    rows=[
        ["สถานะ", "new / acknowledged / in_progress / referred / closed"],
        ["ระดับ", "warning / urgent / critical"],
        ["สังกัด → อำเภอ → โรงเรียน", "Cascade dropdown"],
        ["ประเภทแบบประเมิน", "ST-5 / PHQ-A / CDI"],
        ["ชั้นเรียน / เพศ", "ตามรายการในระบบ"],
        ["ช่วงวันที่", "วันเริ่ม–สิ้นสุด"],
    ],
    col_widths=[5.5, 10.5],
)

add_screenshot_box(doc, "หน้า Alerts — แสดงตารางรายการ Alert พร้อม Badge สถานะและระดับ")

add_heading(doc, "5.3 การอัปเดตสถานะ Alert", level=2)
for step in [
    "ค้นหา Alert ที่ต้องการจัดการ",
    "คลิกปุ่ม 'เปลี่ยนสถานะ' ในแถวนั้น",
    "เลือกสถานะใหม่: Acknowledged / In Progress / Referred / Closed",
    "บันทึกหมายเหตุการดำเนินการ (ถ้ามี)",
    "คลิก 'ยืนยัน'",
]:
    add_numbered(doc, step)

add_screenshot_box(doc, "Modal อัปเดตสถานะ Alert — แสดงตัวเลือกสถานะและช่องหมายเหตุ")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 6: Reports
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 6: รายงานรายบุคคล", level=1)
add_para(doc, "เส้นทาง: /admin/reports", color=C_LGRAY, size=12)

add_heading(doc, "6.1 ข้อมูลในตาราง (50 รายการต่อหน้า)", level=2)
add_table(doc,
    headers=["คอลัมน์", "คำอธิบาย"],
    rows=[
        ["วันที่เข้าทำ", "วันและเวลาที่นักเรียนส่งแบบประเมิน"],
        ["ประเภท", "ST-5 / PHQ-A / CDI"],
        ["นักเรียน", "ชื่อ-นามสกุล (masked บางส่วน)"],
        ["ชั้น/ห้อง", "ระดับและห้องเรียน"],
        ["โรงเรียน", "ชื่อโรงเรียน"],
        ["คะแนน", "คะแนนที่ได้"],
        ["ระดับ", "ปกติ / เล็กน้อย / ปานกลาง / รุนแรง"],
        ["⚑ Suicide Risk", "Badge แดงหากพบ suicide_risk = true"],
    ],
    col_widths=[4.5, 11.5],
)

add_screenshot_box(doc, "หน้า Reports — แสดงตารางข้อมูลรายบุคคลพร้อม Badge ระดับความเสี่ยง")

add_heading(doc, "6.2 การส่งออกรายงาน", level=2)
add_table(doc,
    headers=["รูปแบบ", "ไฟล์", "คำอธิบาย"],
    rows=[
        ["📄 PDF", "lemcs_report.pdf", "รายงานสรุปพร้อม Chart"],
        ["📊 Excel", "lemcs_report.xlsx", "ข้อมูลตาราง สำหรับ Excel/SPSS"],
    ],
    col_widths=[3.0, 4.5, 8.5],
)

add_screenshot_box(doc, "ปุ่ม Export PDF และ Excel บนหน้า Reports")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 7: Students
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 7: ข้อมูลนักเรียน", level=1)
add_para(doc, "เส้นทาง: /admin/students", color=C_LGRAY, size=12)

add_heading(doc, "7.1 ข้อมูลในตาราง (20 รายการต่อหน้า)", level=2)
add_table(doc,
    headers=["คอลัมน์", "คำอธิบาย"],
    rows=[
        ["รหัสนักเรียน", "รหัสประจำตัวตามทะเบียนโรงเรียน"],
        ["ชื่อ-นามสกุล", "—"],
        ["เพศ", "👦 ชาย / 👧 หญิง"],
        ["ชั้น/ห้อง", "ระดับและห้องเรียน"],
        ["วันเดือนปีเกิด", "—"],
        ["โรงเรียน", "—"],
        ["สถานะบัญชี", "✅ เปิด / ❌ ปิด"],
        ["วันที่สร้าง", "วันที่นำเข้าข้อมูล"],
    ],
    col_widths=[4.5, 11.5],
)

add_screenshot_box(doc, "หน้า Students — แสดง FilterBar, ตาราง และ Active Filter Chips")

add_heading(doc, "7.2 การค้นหาและกรอง", level=2)
add_table(doc,
    headers=["ตัวกรอง", "รายละเอียด"],
    rows=[
        ["สังกัด → อำเภอ → โรงเรียน", "Cascade dropdown"],
        ["ชั้นเรียน", "ป.1–ป.6, ม.1–ม.6, ปวช./ปวส."],
        ["ห้อง", "พิมพ์ค้นหา"],
        ["เพศ", "ชาย / หญิง"],
        ["สถานะบัญชี", "เปิด / ปิด"],
        ["ค้นหา", "ชื่อ / รหัสนักเรียน / เลขบัตร (delay 350ms)"],
    ],
    col_widths=[5.5, 10.5],
)

add_heading(doc, "7.3 การแก้ไขข้อมูลนักเรียน", level=2)
for step in [
    "คลิกปุ่ม ✏️ แก้ไข ที่แถวนักเรียน",
    "แก้ไขข้อมูลที่ต้องการ (ชื่อ / ชั้น / ห้อง / เพศ / วันเกิด)",
    "คลิก 'บันทึก'",
]:
    add_numbered(doc, step)
add_note_box(doc, "📌 หมายเหตุ: รหัสนักเรียนไม่สามารถแก้ไขได้หลังจากสร้างแล้ว")

add_screenshot_box(doc, "Modal แก้ไขข้อมูลนักเรียน — แสดงช่องแก้ไขชื่อ/ชั้น/ห้อง")

add_heading(doc, "7.4 การอัปเดตเลขบัตรประชาชน (ข้อมูลสำคัญด้าน PDPA)", level=2)
add_note_box(doc,
    "⚠️ ข้อมูลส่วนบุคคลที่ละเอียดอ่อน — การแก้ไขทุกครั้งจะถูกบันทึกใน Audit Log",
    bg='FFF0E8'
)
for step in [
    "คลิกปุ่ม 🪪 อัปเดต NID ที่แถวนักเรียน",
    "Modal แสดงคำเตือน PDPA — อ่านและทำความเข้าใจ",
    "กรอกเลขบัตรประชาชน: รูปแบบไทย 13 หลัก หรือ G-Code (G + 12 หลัก)",
    "ระบบตรวจสอบ Checksum อัตโนมัติ",
    "คลิก 'ยืนยัน' เพื่อบันทึก",
]:
    add_numbered(doc, step)

add_screenshot_box(doc, "Modal อัปเดต NID — แสดงคำเตือน PDPA และช่องกรอกเลขบัตร")

add_heading(doc, "7.5 การปิด/เปิดบัญชีนักเรียน", level=2)
add_bullet(doc, "ปิดบัญชี (🚫 Deactivate): นักเรียนไม่สามารถ Login ได้ทันที แต่ข้อมูลยังคงอยู่")
add_bullet(doc, "เปิดบัญชี (✅ Activate): คืนสิทธิ์การ Login ให้นักเรียน")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 8: Users
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 8: การจัดการผู้ใช้งาน", level=1)
add_para(doc, "เส้นทาง: /admin/users  |  สิทธิ์: Superadmin และ Systemadmin เท่านั้น", color=C_LGRAY, size=12)

add_heading(doc, "8.1 ข้อมูลในตาราง", level=2)
add_table(doc,
    headers=["คอลัมน์", "คำอธิบาย"],
    rows=[
        ["ชื่อผู้ใช้", "Username สำหรับ Login"],
        ["บทบาท", "Badge แสดง Role"],
        ["องค์กร", "🏫 โรงเรียน / 🏢 อำเภอ / 🏛️ สังกัด / 🌐 ทั้งหมด"],
        ["เข้าสู่ระบบล่าสุด", "วันที่ หรือ 'ยังไม่เคย'"],
        ["สถานะ", "เปิด / ปิด"],
    ],
    col_widths=[4.5, 11.5],
)

add_screenshot_box(doc, "หน้า Users — แสดง Badge สรุป Role และตารางรายชื่อผู้ใช้")

add_heading(doc, "8.2 การสร้างบัญชีผู้ใช้ใหม่", level=2)
for step in [
    "คลิก '➕ เพิ่มผู้ใช้'",
    "กรอก Username (ไม่ซ้ำกับที่มีอยู่)",
    "กรอก Password (ต้องมีอย่างน้อย 8 ตัวอักษร)",
    "เลือก Role จาก Radio Group",
    "กำหนดองค์กร: schooladmin → 3 ขั้น  |  commissionadmin → 2 ขั้น  |  superadmin → ไม่ต้องผูก",
    "คลิก 'สร้างบัญชี'",
]:
    add_numbered(doc, step)
add_note_box(doc, "📌 Superadmin ไม่สามารถสร้างบัญชี systemadmin ได้ (เฉพาะ systemadmin เท่านั้น)")

add_screenshot_box(doc, "Modal สร้างบัญชีใหม่ — แสดงช่องกรอก Username/Password และการเลือก Role/องค์กร")

add_heading(doc, "8.3 การ Reset Password", level=2)
for step in [
    "คลิกปุ่ม 🔑 Reset Password ที่แถวผู้ใช้",
    "กรอก Password ใหม่ (ต้องมีอย่างน้อย 8 ตัวอักษร)",
    "ยืนยันอีกครั้ง",
    "คลิก 'ยืนยัน'",
]:
    add_numbered(doc, step)
add_note_box(doc, "⚡ Password เก่าจะใช้ไม่ได้ทันที — ผู้ใช้ต้อง Login ด้วย Password ใหม่ในครั้งถัดไป")

add_heading(doc, "8.4 การปิด/เปิดบัญชีผู้ใช้", level=2)
add_bullet(doc, "ปิดบัญชี: ผู้ใช้ไม่สามารถ Login ได้ทันที (Soft Delete — ข้อมูลยังอยู่)")
add_bullet(doc, "เปิดบัญชี: คืนสิทธิ์การ Login")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 9: Schools
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 9: ข้อมูลโรงเรียน", level=1)
add_para(doc, "เส้นทาง: /admin/schools", color=C_LGRAY, size=12)
add_note_box(doc, "📌 Superadmin สามารถดูและ Export ข้อมูลโรงเรียนได้ แต่ไม่สามารถเพิ่ม/แก้ไข/ลบได้ (เฉพาะ systemadmin)")

add_heading(doc, "9.1 Summary Cards", level=2)
for item in [
    "จำนวนโรงเรียนทั้งหมด",
    "โรงเรียนที่นำเข้าข้อมูลแล้ว",
    "โรงเรียนที่ยังไม่มีข้อมูล",
    "จำนวนนักเรียนทั้งหมด",
]:
    add_bullet(doc, item)

add_heading(doc, "9.2 ตาราง", level=2)
add_table(doc,
    headers=["คอลัมน์", "คำอธิบาย"],
    rows=[
        ["ชื่อโรงเรียน", "—"],
        ["สังกัด", "สพฐ. / สพม. / สอศ. / สกร."],
        ["อำเภอ", "—"],
        ["ประเภท", "ประถม / มัธยม / อาชีวะ / เอกชน / สกร."],
        ["จำนวนนักเรียน", "—"],
        ["นำเข้าล่าสุด", "วันที่ Import ครั้งล่าสุด"],
        ["สถานะ", "🟢 ล่าสุด (<30 วัน) / 🟡 นานแล้ว / ⚫ ยังไม่มีข้อมูล"],
    ],
    col_widths=[4.5, 11.5],
)

add_screenshot_box(doc, "หน้า Schools — แสดง Summary Cards และตารางโรงเรียน")

add_heading(doc, "9.3 การส่งออก", level=2)
add_bullet(doc, "📊 Excel — ส่งออกตารางโรงเรียนทั้งหมด")
add_bullet(doc, "📄 PDF — รายงานสรุปพร้อม Summary Stats")
add_bullet(doc, "📥 CSV — ส่งออกเป็นไฟล์ CSV")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 10: Audit Logs
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 10: Audit Logs", level=1)
add_para(doc, "เส้นทาง: /admin/audit-logs  |  สิทธิ์: Superadmin และ Systemadmin เท่านั้น", color=C_LGRAY, size=12)
add_para(doc,
    "Audit Log คือบันทึกการเข้าถึงข้อมูลส่วนบุคคล (PII) ที่จำเป็นต้องมีตาม "
    "พระราชบัญญัติคุ้มครองข้อมูลส่วนบุคคล (PDPA) พ.ศ. 2562"
)

add_heading(doc, "10.1 Actions ที่ถูกบันทึก", level=2)
add_table(doc,
    headers=["Action", "ความหมาย"],
    rows=[
        ["view_alert", "การเปิดดูรายละเอียด Alert"],
        ["update_alert", "การเปลี่ยนสถานะ Alert"],
        ["update_national_id", "การแก้ไขเลขบัตรประชาชนนักเรียน"],
    ],
    col_widths=[5.5, 10.5],
)

add_screenshot_box(doc, "หน้า Audit Logs — แสดงตารางบันทึกพร้อม Filter และ Expandable Details")

add_heading(doc, "10.2 การกรอง Audit Logs", level=2)
add_bullet(doc, "Action Type: เลือกประเภทการกระทำ")
add_bullet(doc, "Resource Type: alert / student")
add_bullet(doc, "ช่วงวันที่: วันเริ่ม–สิ้นสุด")
add_para(doc, "แสดง 50 รายการต่อหน้า พร้อม Pagination")

add_heading(doc, "10.3 ความสำคัญทางกฎหมาย", level=2)
add_para(doc,
    "การมี Audit Log แสดงให้เห็นว่าระบบมีความโปร่งใสในการเข้าถึงข้อมูลส่วนบุคคล "
    "ตามมาตรา 36 ของ พ.ร.บ. คุ้มครองข้อมูลส่วนบุคคล พ.ศ. 2562"
)
add_note_box(doc, "📋 แนะนำ: ตรวจสอบ Audit Log เป็นประจำทุกเดือนเพื่อให้แน่ใจว่าการเข้าถึงข้อมูลเป็นไปตามที่ได้รับอนุญาต")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 11: Settings
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 11: การตั้งค่าระบบ", level=1)
add_para(doc, "เส้นทาง: /admin/settings", color=C_LGRAY, size=12)

add_heading(doc, "11.1 ข้อมูลระบบ", level=2)
add_table(doc,
    headers=["รายการ", "ค่า"],
    rows=[
        ["ชื่อระบบ", "LEMCS — Loei Educational MindCare System"],
        ["เวอร์ชัน", "1.0.0"],
        ["ผู้ดูแล", "สำนักงานศึกษาธิการจังหวัดเลย"],
        ["Tech Stack", "Next.js 14 + FastAPI + PostgreSQL"],
    ],
    col_widths=[4.5, 11.5],
)

add_heading(doc, "11.2 Danger Zone (⚠️ ระวัง — ข้อมูลจะหายถาวร)", level=2)
add_note_box(doc,
    "🚨 คำเตือน: การดำเนินการในส่วนนี้จะลบข้อมูลถาวร ไม่สามารถกู้คืนได้\n"
    "ใช้สำหรับล้างข้อมูลทดสอบหรือรีเซ็ตข้อมูลก่อนเริ่มปีการศึกษาใหม่เท่านั้น",
    bg='FFE8E8'
)
add_para(doc, "ลบข้อมูลนักเรียนตามสังกัด (Truncate By Affiliation):", bold=True)
for step in [
    "เลือกสังกัดจาก Dropdown",
    "Modal แสดงชื่อสังกัดที่จะลบ พร้อมจำนวนนักเรียน",
    "ยืนยัน → ลบนักเรียนทั้งหมดในสังกัดนั้น (ไม่สามารถกู้คืนได้)",
]:
    add_numbered(doc, step)

add_screenshot_box(doc, "หน้า Settings — แสดง System Info และ Danger Zone")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# CHAPTER 12: Workflows
# ════════════════════════════════════════════════════════════════
add_heading(doc, "บทที่ 12: ขั้นตอนการทำงานทั่วไป", level=1)

workflows = [
    ("Workflow 1: เริ่มต้นปีการศึกษาใหม่", [
        ("ตรวจสอบโรงเรียน", "/admin/schools → ยืนยันว่าโรงเรียนทั้งหมดมีอยู่ในระบบ"),
        ("ตรวจสอบบัญชีผู้ใช้", "/admin/users → ยืนยันว่า schooladmin ทุกโรงเรียนบัญชีเปิดอยู่"),
        ("เปิดรอบสำรวจใหม่", "/admin/survey-rounds → กด 'เปิดรอบใหม่' กรอก label เช่น 'ภาคเรียน 1/2568'"),
        ("แจ้งโรงเรียน", "แจ้งให้ schooladmin แต่ละโรงเรียนทราบว่าเปิดรอบแล้ว"),
        ("ติดตาม Dashboard", "/admin/dashboard → ดูความคืบหน้า (auto-refresh 60s)"),
    ]),
    ("Workflow 2: ติดตามนักเรียนความเสี่ยงสูง", [
        ("ตรวจสอบ Alerts", "/admin/alerts → กรอง 'Critical' หรือ 'รอดำเนินการ'"),
        ("ดูรายละเอียด", "คลิก 'ดูประวัตินักเรียน' → ดูผลการประเมินทั้งหมด"),
        ("อัปเดตสถานะ", "เปลี่ยนสถานะ → 'In Progress' หรือ 'Referred' บันทึกหมายเหตุ"),
        ("ติดตามผล", "ตรวจสอบซ้ำจนสถานะเป็น 'Closed'"),
    ]),
    ("Workflow 3: ส่งออกรายงานให้ผู้บริหาร", [
        ("เปิด Dashboard", "/admin/dashboard"),
        ("ตั้งค่า Filter", "เลือก รอบสำรวจ + สังกัด + ช่วงวันที่ที่ต้องการ"),
        ("Export", "คลิก '📊 Excel' หรือ '📄 PDF' — ไฟล์ดาวน์โหลดอัตโนมัติ"),
        ("รายงานรายบุคคล", "/admin/reports → กรองเพิ่มเติม → Export (ถ้าต้องการ)"),
    ]),
    ("Workflow 4: แก้ไขข้อมูลนักเรียนที่ผิดพลาด", [
        ("ค้นหานักเรียน", "/admin/students → พิมพ์ชื่อหรือรหัสนักเรียน"),
        ("แก้ไขข้อมูล", "คลิก ✏️ แก้ไข → แก้ไขข้อมูลที่ต้องการ → บันทึก"),
        ("อัปเดต NID", "คลิก 🪪 → อ่านคำเตือน PDPA → กรอกเลขบัตรใหม่ → ยืนยัน (ถ้าจำเป็น)"),
        ("ตรวจสอบ Audit Log", "/admin/audit-logs → ยืนยันว่าบันทึกการแก้ไขอยู่ครบ"),
    ]),
    ("Workflow 5: จัดการบัญชีโรงเรียนที่มีปัญหา Login", [
        ("ตรวจสอบสถานะบัญชี", "/admin/users → ค้นหา Username → ดูว่าสถานะ 'ปิด' หรือไม่"),
        ("เปิดบัญชี", "คลิก ✅ Activate (ถ้าบัญชีปิดอยู่)"),
        ("Reset Password", "คลิก 🔑 → กรอก Password ใหม่ → ยืนยัน → แจ้ง schooladmin (ถ้าลืมรหัสผ่าน)"),
    ]),
]

for wf_title, steps in workflows:
    add_heading(doc, wf_title, level=2)
    for i, (step_name, detail) in enumerate(steps, 1):
        p = doc.add_paragraph()
        r_num = p.add_run(f"ขั้น {i}  ")
        r_num.font.name = 'TH Sarabun New'
        r_num.font.size = Pt(13)
        r_num.font.bold = True
        r_num.font.color.rgb = C_BLUE
        r_bold = p.add_run(f"{step_name}: ")
        r_bold.font.name = 'TH Sarabun New'
        r_bold.font.size = Pt(13)
        r_bold.font.bold = True
        r_detail = p.add_run(detail)
        r_detail.font.name = 'TH Sarabun New'
        r_detail.font.size = Pt(13)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_screenshot_box(doc, "หน้า Alerts กรอง Critical — แสดงรายการที่ต้องดำเนินการ (ตัวอย่าง Workflow 2)")

add_page_break(doc)

# ════════════════════════════════════════════════════════════════
# APPENDIX: FAQ
# ════════════════════════════════════════════════════════════════
add_heading(doc, "ภาคผนวก: คำถามที่พบบ่อย", level=1)

faqs = [
    ("นักเรียนบอกว่าทำแบบประเมินไม่ได้ ทั้งที่ Login ได้แล้ว",
     "ตรวจสอบว่ามีรอบสำรวจที่เปิดอยู่หรือไม่ที่ /admin/survey-rounds ถ้าไม่มีรอบที่สถานะ 'เปิด' นักเรียนจะเห็น Banner สีเหลือง 'ยังไม่เปิดรอบการสำรวจ'"),
    ("ลืมจะปิดรอบสำรวจ นักเรียนยังส่งข้อมูลมาเรื่อยๆ",
     "ไปที่ /admin/survey-rounds → คลิก 'ปิดรอบ' หรือ '⏸ ยกเลิกรอบ' ได้ทันที ข้อมูลที่ส่งมาแล้วจะไม่หาย"),
    ("Alert ขึ้นเยอะมาก จะกรองเฉพาะที่สำคัญได้อย่างไร",
     "ที่ /admin/alerts กรองสถานะ = 'new' และระดับ = 'critical' จะเห็นเฉพาะเคสใหม่ที่ยังไม่มีผู้รับและมีความเสี่ยงสูงสุด"),
    ("ส่งออก Excel แล้วไม่มีข้อมูล",
     "ตรวจสอบว่า FilterBar ไม่ได้กรองข้อมูลจนเหลือ 0 รายการ ลองกด 'ล้างตัวกรอง' แล้ว Export ใหม่"),
    ("โรงเรียนใหม่ยังไม่อยู่ในระบบ จะทำอย่างไร",
     "Superadmin ไม่สามารถเพิ่มโรงเรียนได้โดยตรง ให้ติดต่อ systemadmin (ทีมพัฒนา) เพื่อเพิ่มโรงเรียน"),
    ("จะรู้ได้อย่างไรว่ามีนักเรียนเสี่ยงใหม่เข้ามา",
     "Section B บน Dashboard refresh อัตโนมัติทุก 30 วินาที Alert ใหม่จะมี Badge กะพริบ นอกจากนี้ยังดูได้ที่ /admin/alerts โดยกรองสถานะ = 'new'"),
    ("อยากดูข้อมูลเฉพาะโรงเรียนหนึ่ง จะทำอย่างไร",
     "ใช้ FilterBar บน Dashboard เลือก สังกัด → อำเภอ → โรงเรียน ข้อมูลทุก Section จะอัปเดตพร้อมกัน"),
    ("ต้องการเปรียบเทียบผลการประเมินระหว่างอำเภอ",
     "ที่ Section D บน Dashboard เลือกสังกัดที่ต้องการ ระบบจะ Auto-group by อำเภอ แสดง OrgCompareChart เปรียบเทียบ % ได้ทันที"),
]

for i, (q, a) in enumerate(faqs, 1):
    qp = doc.add_paragraph()
    qr = qp.add_run(f"Q{i}: {q}")
    qr.font.name = 'TH Sarabun New'
    qr.font.size = Pt(13.5)
    qr.font.bold = True
    qr.font.color.rgb = C_BLUE
    qp.paragraph_format.space_before = Pt(8)
    qp.paragraph_format.space_after  = Pt(2)

    ap = doc.add_paragraph()
    ar = ap.add_run(f"A: {a}")
    ar.font.name = 'TH Sarabun New'
    ar.font.size = Pt(13)
    ar.font.color.rgb = C_DGRAY
    ap.paragraph_format.left_indent  = Cm(0.5)
    ap.paragraph_format.space_after  = Pt(4)

add_divider(doc)

# Footer
foot_p = doc.add_paragraph()
foot_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
foot_r = foot_p.add_run(
    "คู่มือนี้จัดทำสำหรับ LEMCS เวอร์ชัน 1.0  ·  "
    "หากระบบมีการอัปเดตโปรดตรวจสอบเวอร์ชันล่าสุดจากทีมพัฒนา\n"
    "พัฒนาเพื่อนักเรียนจังหวัดเลยกว่า 100,000 คน  ·  พ.ศ. 2568"
)
foot_r.font.name = 'TH Sarabun New'
foot_r.font.size = Pt(11)
foot_r.font.color.rgb = C_LGRAY
foot_p.paragraph_format.space_before = Pt(16)

# ════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"[OK] Saved: {OUT}")
print(f"     Pages: ~{len(doc.paragraphs) // 35 + 1} (estimated)")
