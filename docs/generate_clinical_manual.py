"""
สร้างเอกสาร Word คู่มือตรวจสอบเครื่องมือประเมินสุขภาพจิต LEMCS
สำหรับจิตแพทย์/นักจิตวิทยาตรวจสอบความถูกต้องทางคลินิก
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ========== สีหลัก ==========
COLOR_TITLE      = RGBColor(0x1e, 0x3a, 0x5f)   # navy blue
COLOR_H1         = RGBColor(0x1e, 0x3a, 0x5f)
COLOR_H2         = RGBColor(0x1a, 0x56, 0x76)   # teal-ish
COLOR_H3         = RGBColor(0x2d, 0x6a, 0x4f)   # green
COLOR_TABLE_HDR  = RGBColor(0x1e, 0x3a, 0x5f)
COLOR_ALT_ROW    = RGBColor(0xf0, 0xf4, 0xf8)
COLOR_RED_ROW    = RGBColor(0xff, 0xeb, 0xee)
COLOR_YELLOW_ROW = RGBColor(0xff, 0xf9, 0xe6)
COLOR_GREEN_ROW  = RGBColor(0xe8, 0xf5, 0xe9)
COLOR_BLUE_ROW   = RGBColor(0xe3, 0xf2, 0xfd)
COLOR_WARN_BOX   = RGBColor(0xff, 0xf3, 0xe0)
COLOR_INFO_BOX   = RGBColor(0xe8, 0xf4, 0xfd)
COLOR_CODE_BG    = RGBColor(0xf5, 0xf5, 0xf5)
WHITE            = RGBColor(0xff, 0xff, 0xff)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), val.get("val", "single"))
            el.set(qn("w:sz"), val.get("sz", "4"))
            el.set(qn("w:color"), val.get("color", "AAAAAA"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_paragraph_shading(paragraph, rgb: RGBColor):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}")
    pPr.append(shd)


def set_page_margins(doc, top=2.0, bottom=2.0, left=2.5, right=2.0):
    section = doc.sections[0]
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = COLOR_TITLE
    run.font.name = "TH Sarabun New"
    p.space_after = Pt(4)
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.name = "TH Sarabun New"
    p.space_after = Pt(2)
    return p


def add_h1(doc, text, number=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"{'ส่วนที่ '+number+' — ' if number else ''}{text}")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = COLOR_H1
    run.font.name = "TH Sarabun New"
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_H2
    run.font.name = "TH Sarabun New"
    return p


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_H3
    run.font.name = "TH Sarabun New"
    return p


def add_body(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.bold = bold
    run.font.name = "TH Sarabun New"
    return p


def add_bullet(doc, text, level=0, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(12)
        rb.font.name = "TH Sarabun New"
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "TH Sarabun New"
    return p


def add_note_box(doc, text, box_color=None, prefix=""):
    if box_color is None:
        box_color = COLOR_INFO_BOX
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    add_paragraph_shading(p, box_color)
    if prefix:
        rb = p.add_run(prefix + " ")
        rb.bold = True
        rb.font.size = Pt(12)
        rb.font.name = "TH Sarabun New"
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = "TH Sarabun New"
    return p


def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    add_paragraph_shading(p, COLOR_CODE_BG)
    p.paragraph_format.left_indent  = Cm(0.7)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    for line in code_text.strip().split("\n"):
        if p.runs:
            p.add_run("\n")
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x37, 0x47, 0x4f)
    return p


def make_table(doc, headers, rows, col_widths=None,
               row_colors=None, header_bg=None, font_size=11):
    """
    row_colors: list ความยาวเท่า rows แต่ละ element เป็น RGBColor หรือ None
    """
    if header_bg is None:
        header_bg = COLOR_TABLE_HDR
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ส่วนหัว
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(font_size)
        hdr_cells[i].paragraphs[0].runs[0].font.name = "TH Sarabun New"
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], header_bg)

    # แถวข้อมูล
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = None
        if row_colors and r_idx < len(row_colors):
            bg = row_colors[r_idx]
        elif r_idx % 2 == 1:
            bg = COLOR_ALT_ROW
        for c_idx, cell_text in enumerate(row):
            # รองรับ tuple (text, bold)
            if isinstance(cell_text, tuple):
                text, is_bold = cell_text
            else:
                text, is_bold = str(cell_text), False
            row_cells[c_idx].text = text
            run = row_cells[c_idx].paragraphs[0].runs[0] if row_cells[c_idx].paragraphs[0].runs else row_cells[c_idx].paragraphs[0].add_run(text)
            run.font.size = Pt(font_size)
            run.font.name = "TH Sarabun New"
            run.bold = is_bold
            if bg:
                set_cell_bg(row_cells[c_idx], bg)

    # กำหนดความกว้าง
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "BBBBBB")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ==========================================
#  สร้างเอกสาร
# ==========================================
doc = Document()
set_page_margins(doc)

# ---- ปกหัวเรื่อง ----
doc.add_paragraph()
add_title(doc, "คู่มือตรวจสอบเครื่องมือประเมินสุขภาพจิต")
add_title(doc, "ระบบ LEMCS")
add_subtitle(doc, "Loei Educational MindCare System")
doc.add_paragraph()
add_subtitle(doc, "สำหรับจิตแพทย์ / นักจิตวิทยา ตรวจสอบความถูกต้องทางคลินิก")
add_divider(doc)
add_subtitle(doc, "วันที่จัดทำ: 13 เมษายน พ.ศ. 2569   |   ระบบรองรับนักเรียน 100,000+ คน จังหวัดเลย")
doc.add_paragraph()

# ==========================================
#  ส่วนที่ 1 — ภาพรวม
# ==========================================
add_h1(doc, "ภาพรวมระบบและแบบประเมิน", "1")
add_body(doc,
    "LEMCS ประเมินสุขภาพจิตนักเรียน K–12 โดยใช้แบบประเมิน 3 ชุด "
    "ซึ่งระบบเลือกแบบประเมินซึมเศร้าให้อัตโนมัติตามอายุของนักเรียน:")

make_table(doc,
    ["แบบประเมิน", "กลุ่มเป้าหมาย (อายุ)", "วัตถุประสงค์", "จำนวนข้อ", "ช่วงคะแนน"],
    [
        ["ST-5",  "นักเรียนทุกคน",      "ประเมินความเครียด",                     "5 ข้อ",  "0–15"],
        ["PHQ-A", "อายุ 11–20 ปี",      "ประเมินภาวะซึมเศร้าวัยรุ่น",            "9+2 ข้อ","0–27"],
        ["CDI",   "อายุต่ำกว่า 11 ปี",  "ประเมินภาวะซึมเศร้าในเด็ก",             "27 ข้อ", "0–54"],
    ],
    col_widths=[3.0, 3.5, 5.5, 2.5, 2.5],
    row_colors=[COLOR_BLUE_ROW, COLOR_YELLOW_ROW, COLOR_GREEN_ROW],
)

add_note_box(doc,
    "กฎการเลือกแบบประเมิน: หากอายุ 7–17 ปี → ระบบให้ทำ CDI  |  หากอายุ ≥ 18 ปี → ระบบให้ทำ PHQ-A  |  "
    "หากอายุ < 7 ปี → ทำ ST-5 เท่านั้น  |  นักเรียนทุกคนอายุ 7 ปีขึ้นไปทำ ST-5 ได้เสมอ",
    COLOR_INFO_BOX, "ℹ️")

# ==========================================
#  ส่วนที่ 2 — Workflow นักเรียน
# ==========================================
add_h1(doc, "ขั้นตอนการใช้งานของนักเรียน", "2")

add_h2(doc, "2.1  การเข้าสู่ระบบ")
add_body(doc, "นักเรียนเข้าสู่ระบบด้วยข้อมูล 3 รายการพร้อมกัน:")
add_bullet(doc, "รหัสนักเรียน (Student Code)")
add_bullet(doc, "วันเกิด — ระบบใช้คำนวณอายุเพื่อเลือกแบบประเมินที่เหมาะสม")
add_bullet(doc, "เลขบัตรประจำตัวประชาชน — เข้ารหัส AES-256 ตาม พ.ร.บ. PDPA")

add_h2(doc, "2.2  หน้าจอยินยอม PDPA (ครั้งแรกเท่านั้น)")
add_body(doc, "ก่อนทำแบบประเมินครั้งแรก นักเรียนต้องอ่านและยืนยันยินยอมก่อน "
              "ระบบถึงจะดำเนินการต่อได้ โดยมีปุ่ม \"ไม่ยินยอม\" เพื่อออกจากระบบ")

add_h2(doc, "2.3  ขั้นตอนการทำแบบประเมิน")
add_bullet(doc, "แสดงทีละ 1 ข้อ พร้อมแถบความก้าวหน้า (Progress Bar)")
add_bullet(doc, "เลือกคำตอบ → ระบบเลื่อนไปข้อถัดไปอัตโนมัติ")
add_bullet(doc, "มีปุ่มย้อนกลับ เพื่อแก้ไขคำตอบข้อก่อนหน้าได้")
add_bullet(doc, "บันทึกอัตโนมัติทุก 30 วินาที (กรณีสัญญาณหลุด)")
add_bullet(doc, "ต้องตอบครบทุกข้อจึงจะส่งได้ — มีหน้าต่างยืนยันก่อนส่ง")
add_note_box(doc, "เมื่อส่งแล้วไม่สามารถแก้ไขคำตอบได้", COLOR_WARN_BOX, "⚠️")

add_h2(doc, "2.4  หน้าแสดงผล")
add_bullet(doc, "คะแนนตัวเลข (แสดงขนาดใหญ่)")
add_bullet(doc, "ป้ายระดับความรุนแรง (ระบายสีตามความเสี่ยง)")
add_bullet(doc, "คำแนะนำเบื้องต้น (ขึ้นกับระดับที่ได้)")
add_bullet(doc, "ปุ่มโทรสายด่วนสุขภาพจิต 1323 / 1387 (กรณีความเสี่ยงสูง)")

add_divider(doc)

# ==========================================
#  ส่วนที่ 3 — ST-5
# ==========================================
add_h1(doc, "แบบประเมินความเครียด ST-5", "3")

add_h2(doc, "3.1  ที่มาและมาตรฐานอ้างอิง")
add_body(doc, "ST-5 (Stress Test 5) พัฒนาโดยกรมสุขภาพจิต กระทรวงสาธารณสุข ประเทศไทย "
              "ใช้คัดกรองความเครียดในประชากรทั่วไป รวมถึงนักเรียน")

add_h2(doc, "3.2  ข้อคำถาม")
make_table(doc,
    ["ข้อ", "คำถาม", "ตัวเลือกและคะแนน"],
    [
        ["Q1", "คุณรู้สึกไม่สบายใจ กังวล หรือเครียดบ้างไหม?",
         "ไม่เลย(0) / บางครั้ง(1) / บ่อยครั้ง(2) / เกือบทุกวัน(3)"],
        ["Q2", "คุณรู้สึกหงุดหงิด รำคาญใจ หรือโกรธง่ายบ้างไหม?",
         "ไม่เลย(0) / บางครั้ง(1) / บ่อยครั้ง(2) / เกือบทุกวัน(3)"],
        ["Q3", "คุณรู้สึกเหนื่อยในการใช้ชีวิต หรือรู้สึกว่าตัวเองมีภาระมากเกินไปบ้างไหม?",
         "ไม่เลย(0) / บางครั้ง(1) / บ่อยครั้ง(2) / เกือบทุกวัน(3)"],
        ["Q4", "คุณมีปัญหาการนอนหลับ เช่น นอนไม่หลับ ตื่นกลางดึก หรือหลับมากผิดปกติบ้างไหม?",
         "ไม่เลย(0) / บางครั้ง(1) / บ่อยครั้ง(2) / เกือบทุกวัน(3)"],
        ["Q5", "คุณรู้สึกว่าความเครียดส่งผลต่อการเรียน การใช้ชีวิต หรือความสัมพันธ์กับคนรอบข้างบ้างไหม?",
         "ไม่เลย(0) / บางครั้ง(1) / บ่อยครั้ง(2) / เกือบทุกวัน(3)"],
    ],
    col_widths=[1.2, 9.5, 6.3],
)

add_h2(doc, "3.3  การคำนวณและเกณฑ์แปลผล")
add_body(doc, "คะแนนรวม = Q1 + Q2 + Q3 + Q4 + Q5   (ช่วง 0–15)")

make_table(doc,
    ["คะแนน", "ระดับ", "ป้ายที่แสดงบนหน้าจอ", "สัญลักษณ์สี"],
    [
        ["0–4",   "normal",   "ปกติ",            "เขียว"],
        ["5–7",   "mild",     "เครียดเล็กน้อย",  "น้ำเงิน"],
        ["8–11",  "moderate", "เครียดปานกลาง",   "เหลือง"],
        ["12–15", "severe",   "เครียดสูง",        "แดง"],
    ],
    col_widths=[2.0, 2.5, 5.0, 3.0],
    row_colors=[COLOR_GREEN_ROW, COLOR_BLUE_ROW, COLOR_YELLOW_ROW, COLOR_RED_ROW],
)

add_note_box(doc, "ST-5 ไม่มีการตรวจความเสี่ยงฆ่าตัวตาย (Suicide Risk = ไม่มีเสมอ) "
                  "เนื่องจากเป็นแบบประเมินความเครียด ไม่ใช่ภาวะซึมเศร้า",
             COLOR_INFO_BOX, "ℹ️")

add_h2(doc, "3.4  คำแนะนำที่แสดงต่อนักเรียน")
make_table(doc,
    ["ระดับ", "คำแนะนำที่ระบบแสดง"],
    [
        ["ปกติ (0–4)",          "คุณมีระดับความเครียดอยู่ในเกณฑ์ปกติ  |  ดูแลสุขภาพจิตต่อไป ออกกำลังกาย พักผ่อนให้เพียงพอ"],
        ["เล็กน้อย (5–7)",      "คุณมีความเครียดเล็กน้อย ไม่ต้องกังวล  |  ลองพักผ่อน ทำกิจกรรมที่ชอบ หรือคุยกับเพื่อนสนิท"],
        ["ปานกลาง (8–11)",      "คุณมีระดับความเครียดปานกลาง  |  แนะนำให้พูดคุยกับครูแนะแนวหรือผู้ปกครอง  |  ฝึกหายใจลึกๆ และจัดการเวลาให้ดีขึ้น"],
        ["สูง (12–15)",         "คุณมีระดับความเครียดสูง  ⚠️ ควรพบครูแนะแนวโดยเร็วที่สุด  |  หากรู้สึกไม่ไหว โทรสายด่วนสุขภาพจิต 1323"],
    ],
    col_widths=[3.5, 13.5],
    row_colors=[COLOR_GREEN_ROW, COLOR_BLUE_ROW, COLOR_YELLOW_ROW, COLOR_RED_ROW],
)

add_divider(doc)

# ==========================================
#  ส่วนที่ 4 — PHQ-A
# ==========================================
add_h1(doc, "แบบประเมินภาวะซึมเศร้าวัยรุ่น PHQ-A", "4")

add_h2(doc, "4.1  ที่มาและมาตรฐานอ้างอิง")
add_body(doc,
    "PHQ-A (Patient Health Questionnaire – Adolescent) "
    "ฉบับมาตรฐานของสถาบันสุขภาพจิตเด็กและวัยรุ่นราชนครินทร์ พ.ศ. 2561 "
    "ดัดแปลงจาก PHQ-9 สำหรับวัยรุ่น  กลุ่มเป้าหมาย: อายุ 11–20 ปี")

add_h2(doc, "4.2  ข้อคำถามหลัก Q1–Q9")
add_body(doc, "ตัวเลือกทุกข้อ: ไม่มีเลย = 0  /  มีบางวัน = 1  /  มีมากกว่า 7 วัน = 2  /  มีแทบทุกวัน = 3")

make_table(doc,
    ["ข้อ", "คำถาม", "มิติทางคลินิก", "หมายเหตุ"],
    [
        ["Q1", "รู้สึกซึมเศร้า หงุดหงิด หรือสิ้นหวัง",
         "Depressed mood / Irritability", ""],
        ["Q2", "เบื่อ ไม่ค่อยสนใจหรือไม่เพลิดเพลินเวลาทำสิ่งต่างๆ",
         "Anhedonia", ""],
        ["Q3", "นอนหลับยาก หรือหลับๆ ตื่นๆ หรือหลับมากเกินไป",
         "Sleep disturbance", ""],
        ["Q4", "รู้สึกเหนื่อยล้า หรือไม่ค่อยมีพลัง",
         "Fatigue / Energy loss", ""],
        ["Q5", "ไม่อยากอาหาร น้ำหนักลด หรือกินมากกว่าปกติ",
         "Appetite change", ""],
        ["Q6", "รู้สึกแย่กับตัวเอง หรือรู้สึกว่าตัวเองล้มเหลว หรือทำให้ตัวเองหรือครอบครัวผิดหวัง",
         "Worthlessness / Guilt", ""],
        ["Q7", "จดจ่อกับสิ่งต่างๆ ได้ยาก เช่น ทำการบ้าน อ่านหนังสือ หรือดูโทรทัศน์",
         "Concentration", ""],
        ["Q8", "พูดหรือทำอะไรช้าลงจนคนอื่นสังเกตเห็นได้ หรือกระสับกระส่ายจนต้องเคลื่อนไหวมากกว่าปกติ",
         "Psychomotor changes", ""],
        ["Q9", "คิดว่าถ้าตายไปเสียจะดีกว่า หรือคิดจะทำร้ายตัวเองด้วยวิธีใดวิธีหนึ่ง",
         "Suicidal ideation", "⚠️ trigger Suicide Risk ถ้า ≥ 1"],
    ],
    col_widths=[1.2, 7.5, 4.0, 4.3],
    row_colors=[None, None, None, None, None, None, None, None, COLOR_RED_ROW],
)

add_h2(doc, "4.3  คำถามเพิ่มเติม BQ1–BQ2 (ไม่นับคะแนน — ตรวจ Suicide Risk เท่านั้น)")
make_table(doc,
    ["ข้อ", "คำถาม", "ตัวเลือก", "ผลถ้าตอบว่า \"ใช่\""],
    [
        ["BQ1",
         "ใน 1 เดือนที่ผ่านมา มีช่วงไหนที่คุณมีความคิดอยากตาย หรือไม่อยากมีชีวิตอยู่อย่างจริงจังหรือไม่?",
         "ไม่มี / มี",
         "⚠️ Suicide Risk = ใช่"],
        ["BQ2",
         "ตลอดชีวิตที่ผ่านมา คุณเคยพยายามที่จะทำให้ตัวเองตาย หรือลงมือฆ่าตัวตายหรือไม่?",
         "ไม่เคย / เคย",
         "⚠️ Suicide Risk = ใช่"],
    ],
    col_widths=[1.2, 7.5, 2.5, 5.8],
    row_colors=[COLOR_RED_ROW, COLOR_RED_ROW],
)

add_note_box(doc,
    "BQ1 และ BQ2 ไม่ถูกนำมาคิดคะแนนรวม แต่ถ้าตอบ \"มี\" หรือ \"เคย\" "
    "ระบบจะถือว่ามีความเสี่ยงฆ่าตัวตายทันที แม้คะแนนรวม Q1–Q9 จะอยู่ในระดับ \"ไม่มีภาวะซึมเศร้า\" ก็ตาม",
    COLOR_WARN_BOX, "⚠️")

add_h2(doc, "4.4  การคำนวณและเกณฑ์แปลผล")
add_body(doc, "คะแนนรวม = Q1 + Q2 + Q3 + Q4 + Q5 + Q6 + Q7 + Q8 + Q9   (ช่วง 0–27)")

make_table(doc,
    ["คะแนน", "ระดับ", "ป้ายที่แสดงบนหน้าจอ", "สัญลักษณ์สี"],
    [
        ["0–4",   "none",       "ไม่มีภาวะซึมเศร้า", "เขียว"],
        ["5–9",   "mild",       "เล็กน้อย",           "น้ำเงิน"],
        ["10–14", "moderate",   "ปานกลาง",             "เหลือง"],
        ["15–19", "severe",     "มาก",                 "แดง"],
        ["20–27", "very_severe","รุนแรง",              "แดง (เข้ม)"],
    ],
    col_widths=[2.0, 3.0, 5.0, 3.0],
    row_colors=[COLOR_GREEN_ROW, COLOR_BLUE_ROW, COLOR_YELLOW_ROW, COLOR_RED_ROW, COLOR_RED_ROW],
)

add_h2(doc, "4.5  เงื่อนไข Suicide Risk (PHQ-A)")
add_note_box(doc,
    "Suicide Risk = ใช่  เมื่อเข้าเงื่อนไขอย่างน้อย 1 ข้อ:\n"
    "  • Q9 ≥ 1  (ตอบ \"มีบางวัน\" ขึ้นไป)\n"
    "  • BQ1 ตอบ \"มี\"  (มีความคิดอยากตายใน 1 เดือนที่ผ่านมา)\n"
    "  • BQ2 ตอบ \"เคย\"  (เคยพยายามฆ่าตัวตาย)",
    COLOR_RED_ROW, "🚨")

add_divider(doc)

# ==========================================
#  ส่วนที่ 5 — CDI
# ==========================================
add_h1(doc, "แบบประเมินภาวะซึมเศร้าในเด็ก CDI", "5")

add_h2(doc, "5.1  ที่มาและมาตรฐานอ้างอิง")
add_body(doc,
    "CDI (Children's Depression Inventory) พัฒนาโดย Maria Kovacs (1985) ฉบับแปลไทย "
    "ใช้สำหรับเด็กและวัยรุ่นอายุ 7–17 ปี ตามมาตรฐานสากล (MHS CDI 2; ScienceDirect; PMC) "
    "ในระบบนี้ใช้ CDI สำหรับนักเรียนอายุ 7–17 ปี")

add_h2(doc, "5.2  หลักการให้คะแนน (สำคัญมาก)")
add_body(doc,
    "CDI ใช้หลักการให้คะแนนสองแบบ แบ่งข้อคำถามออกเป็น 2 กลุ่ม "
    "บางข้อ \"ก\" ดีที่สุด / ค้ รุนแรงที่สุด  แต่บางข้อ \"ก\" รุนแรงที่สุด / ค ดีที่สุด:")

make_table(doc,
    ["กลุ่ม", "ข้อที่อยู่ในกลุ่ม", "คะแนน ก", "คะแนน ข", "คะแนน ค"],
    [
        ["กลุ่ม A — ตรงทิศทาง\n(ค รุนแรงที่สุด)",
         "1, 3, 4, 6, 9, 12, 14, 17, 19, 20, 22, 23, 26, 27\n(รวม 14 ข้อ)",
         "0", "1", "2"],
        ["กลุ่ม B — กลับทิศทาง\n(ก รุนแรงที่สุด)",
         "2, 5, 7, 8, 10, 11, 13, 15, 16, 18, 21, 24, 25\n(รวม 13 ข้อ)",
         "2", "1", "0"],
    ],
    col_widths=[4.2, 6.5, 1.5, 1.5, 1.5],
    row_colors=[COLOR_BLUE_ROW, COLOR_YELLOW_ROW],
)

add_note_box(doc,
    "คะแนนรวม = ผลรวมคะแนนทั้ง 27 ข้อ (ตามกลุ่ม A หรือ B)   ช่วงคะแนน 0–54",
    COLOR_INFO_BOX, "ℹ️")

add_h2(doc, "5.3  ตารางข้อคำถามทั้ง 27 ข้อ พร้อมคะแนนแต่ละตัวเลือก")
make_table(doc,
    ["ข้อ", "กลุ่ม", "ตัวเลือก ก", "ตัวเลือก ข", "ตัวเลือก ค", "คะแนน ก/ข/ค"],
    [
        ["1",  "A", "ฉันรู้สึกเศร้าบางครั้ง",                          "ฉันรู้สึกเศร้าบ่อยๆ",                            "ฉันรู้สึกเศร้าตลอดเวลา",                      "0/1/2"],
        ["2",  "B", "ไม่มีอะไรดีเลย",                                  "ฉันไม่แน่ใจว่าอะไรจะดีขึ้น",                   "อะไรๆ จะดีขึ้นสำหรับฉัน",                     "2/1/0"],
        ["3",  "A", "ฉันทำสิ่งต่างๆ ได้ดีเกือบทุกอย่าง",               "ฉันทำผิดพลาดหลายอย่าง",                        "ฉันทำทุกอย่างผิดหมด",                         "0/1/2"],
        ["4",  "A", "ฉันสนุกกับหลายสิ่ง",                              "ฉันสนุกกับบางสิ่ง",                             "ไม่มีอะไรสนุกเลย",                             "0/1/2"],
        ["5",  "B", "ฉันเป็นคนเลวตลอดเวลา",                           "ฉันเป็นคนเลวบ่อยครั้ง",                        "ฉันเป็นคนเลวบางครั้ง",                        "2/1/0"],
        ["6",  "A", "ฉันแทบไม่เคยคิดว่าจะมีอะไรร้ายเกิดกับฉัน",      "ฉันกังวลว่าจะมีอะไรร้ายเกิดกับฉัน",           "ฉันแน่ใจว่าจะมีอะไรร้ายเกิดกับฉัน",         "0/1/2"],
        ["7",  "B", "ฉันเกลียดตัวเอง",                                 "ฉันไม่ชอบตัวเอง",                               "ฉันชอบตัวเอง",                                 "2/1/0"],
        ["8",  "B", "สิ่งร้ายๆ ทั้งหมดเป็นความผิดของฉัน",             "สิ่งร้ายๆ หลายอย่างเป็นความผิดของฉัน",        "สิ่งร้ายๆ มักไม่ใช่ความผิดของฉัน",           "2/1/0"],
        ["9 ⚠️","A","ฉันไม่ได้คิดเรื่องตาย",                           "ฉันคิดเรื่องตายแต่ไม่อยากทำ",                  "ฉันอยากตาย",                                   "0/1/2"],
        ["10", "B", "ฉันอยากร้องไห้ทุกวัน",                            "ฉันอยากร้องไห้บ่อยๆ",                          "ฉันอยากร้องไห้เป็นบางครั้ง",                  "2/1/0"],
        ["11", "B", "มีบางอย่างกวนใจฉันตลอดเวลา",                     "มีบางอย่างกวนใจฉันบ่อย",                       "มีบางอย่างกวนใจฉันเป็นบางครั้ง",              "2/1/0"],
        ["12", "A", "ฉันชอบอยู่กับคนอื่น",                             "ฉันไม่อยากอยู่กับคนอื่นบ่อย",                  "ฉันไม่อยากอยู่กับใครเลย",                     "0/1/2"],
        ["13", "B", "ฉันตัดสินใจไม่ได้",                               "ฉันตัดสินใจได้ยาก",                             "ฉันตัดสินใจได้ง่าย",                           "2/1/0"],
        ["14", "A", "ฉันมีรูปร่างหน้าตาดี",                            "มีบางอย่างเกี่ยวกับรูปร่างที่ฉันไม่ชอบ",      "ฉันดูน่าเกลียด",                               "0/1/2"],
        ["15", "B", "ฉันต้องบังคับตัวเองตลอดเวลาให้ทำการบ้าน",       "ฉันต้องบังคับตัวเองบ่อยๆ ให้ทำการบ้าน",       "การทำการบ้านไม่ใช่ปัญหาสำหรับฉัน",           "2/1/0"],
        ["16", "B", "ฉันนอนหลับยากทุกคืน",                             "ฉันนอนหลับยากหลายคืน",                         "ฉันนอนหลับได้ดี",                              "2/1/0"],
        ["17", "A", "ฉันเหนื่อยบางครั้ง",                              "ฉันเหนื่อยบ่อย",                                "ฉันเหนื่อยตลอดเวลา",                           "0/1/2"],
        ["18", "B", "เกือบทุกวันฉันไม่อยากกินอาหาร",                  "หลายวันฉันไม่อยากกินอาหาร",                    "ฉันกินได้ดีปกติ",                              "2/1/0"],
        ["19", "A", "ฉันไม่กังวลเรื่องเจ็บปวดต่างๆ",                  "ฉันกังวลเรื่องเจ็บปวดบ่อย",                    "ฉันกังวลเรื่องเจ็บปวดตลอดเวลา",               "0/1/2"],
        ["20", "A", "ฉันไม่รู้สึกเหงา",                                "ฉันรู้สึกเหงาบ่อย",                             "ฉันรู้สึกเหงาตลอดเวลา",                        "0/1/2"],
        ["21", "B", "ฉันไม่เคยสนุกกับโรงเรียน",                       "ฉันสนุกกับโรงเรียนเป็นบางครั้ง",               "ฉันสนุกกับโรงเรียนบ่อย",                      "2/1/0"],
        ["22", "A", "ฉันมีเพื่อนหลายคน",                               "ฉันมีเพื่อนบ้างแต่อยากมีมากกว่านี้",          "ฉันไม่มีเพื่อนเลย",                            "0/1/2"],
        ["23", "A", "การเรียนของฉันดี",                                 "การเรียนของฉันไม่ดีเหมือนแต่ก่อน",             "ฉันเรียนแย่มากในวิชาที่เคยเรียนได้ดี",        "0/1/2"],
        ["24", "B", "ฉันไม่มีทางเก่งเท่าเด็กคนอื่นๆ",                 "ฉันเก่งเท่าเด็กคนอื่นๆ ถ้าฉันต้องการ",       "ฉันเก่งเท่าเด็กคนอื่นๆ",                      "2/1/0"],
        ["25", "B", "ไม่มีใครรักฉันจริงๆ",                             "ฉันไม่แน่ใจว่ามีใครรักฉัน",                    "ฉันแน่ใจว่ามีคนรักฉัน",                       "2/1/0"],
        ["26", "A", "ฉันทำตามที่คนอื่นบอกเสมอ",                       "ฉันไม่ค่อยทำตามที่คนอื่นบอก",                 "ฉันไม่เคยทำตามที่คนอื่นบอก",                  "0/1/2"],
        ["27", "A", "ฉันเข้ากับคนอื่นได้ดี",                           "ฉันทะเลาะกับคนอื่นบ่อย",                       "ฉันทะเลาะกับคนอื่นตลอดเวลา",                  "0/1/2"],
    ],
    col_widths=[1.2, 1.3, 4.2, 4.2, 4.2, 2.0],
    row_colors=[
        None, None, None, None, None, None, None, None,
        COLOR_RED_ROW,  # Q9
        None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None, None,
    ],
    font_size=10,
)

add_note_box(doc, "ข้อ 9 (แถวสีแดง) — เป็นข้อ Suicide Risk ของ CDI  "
                  "ระบบ trigger Suicide Risk เฉพาะเมื่อเลือก 'ค. ฉันอยากตาย' เท่านั้น  "
                  "การเลือก 'ข. ฉันคิดเรื่องตายแต่ไม่อยากทำ' ยังไม่ถือเป็น Suicide Risk ในระบบนี้",
             COLOR_WARN_BOX, "⚠️")

add_h2(doc, "5.4  การคำนวณและเกณฑ์แปลผล")
make_table(doc,
    ["คะแนน", "ระดับ", "ป้ายที่แสดงบนหน้าจอ", "สัญลักษณ์สี"],
    [
        ["0–14", "normal",   "ปกติ",               "เขียว"],
        ["≥ 15", "clinical", "ควรได้รับการดูแล",   "แดง"],
    ],
    col_widths=[2.0, 2.5, 5.0, 3.0],
    row_colors=[COLOR_GREEN_ROW, COLOR_RED_ROW],
)

add_divider(doc)

# ==========================================
#  ส่วนที่ 6 — Suicide Risk Protocol
# ==========================================
add_h1(doc, "ระบบเตือนภัยวิกฤต (Suicide Risk Protocol)", "6")

add_h2(doc, "6.1  เงื่อนไขกระตุ้น Suicide Risk")
make_table(doc,
    ["แบบประเมิน", "เงื่อนไขที่ทำให้ Suicide Risk = ใช่"],
    [
        ["ST-5",  "ไม่มี — False เสมอ (แบบประเมินความเครียด)"],
        ["PHQ-A", "Q9 ≥ 1  หรือ  BQ1 ตอบ \"มี\"  หรือ  BQ2 ตอบ \"เคย\"  (เข้าอย่างน้อย 1 เงื่อนไข)"],
        ["CDI",   "Q9 ตอบ \"ค. ฉันอยากตาย\" เท่านั้น"],
    ],
    col_widths=[3.0, 14.0],
    row_colors=[COLOR_GREEN_ROW, COLOR_RED_ROW, COLOR_RED_ROW],
)

add_h2(doc, "6.2  สิ่งที่ระบบดำเนินการทันทีเมื่อ Suicide Risk = ใช่")
add_bullet(doc, "แสดงแบนเนอร์เร่งด่วนสีแดง บนหน้าจอนักเรียน พร้อมข้อความ \"เรื่องด่วน ขอความช่วยเหลือทันที\" และปุ่มโทรสายด่วนแบบกดโทรได้ทันที")
add_bullet(doc, "สร้าง Critical Alert ในระบบ → แจ้งเตือนครูแนะแนวและผู้บริหาร ผ่านช่องทาง (ตามลำดับ): LINE Notify → Email → การแจ้งเตือนในแอป")
add_bullet(doc, "การแจ้งเตือนส่งแบบ Synchronous (ไม่ผ่านคิว) เพื่อความรวดเร็วสูงสุด")

add_h2(doc, "6.3  สายด่วนที่แสดงต่อนักเรียน")
make_table(doc,
    ["สายด่วน", "หมายเลข", "บริการ"],
    [
        ["สายด่วนสุขภาพจิต",    "1323", "24 ชั่วโมง / โทรฟรี"],
        ["สายด่วนช่วยเหลือเด็ก","1387", "24 ชั่วโมง"],
        ["กรมสุขภาพจิต",        "02-149-5555", "ช่องทางเพิ่มเติม"],
    ],
    col_widths=[5.0, 3.5, 8.5],
    row_colors=[COLOR_RED_ROW, COLOR_RED_ROW, COLOR_BLUE_ROW],
)

add_divider(doc)

# ==========================================
#  ส่วนที่ 7 — สรุปตาราง
# ==========================================
add_h1(doc, "สรุปเกณฑ์คะแนนทั้งหมด", "7")

add_h2(doc, "7.1  ST-5  (0–15 คะแนน)")
make_table(doc,
    ["คะแนน", "ระดับ", "ป้าย", "Suicide Risk"],
    [
        ["0–4",   "normal",   "ปกติ",            "ไม่มี"],
        ["5–7",   "mild",     "เครียดเล็กน้อย",  "ไม่มี"],
        ["8–11",  "moderate", "เครียดปานกลาง",   "ไม่มี"],
        ["12–15", "severe",   "เครียดสูง",        "ไม่มี"],
    ],
    col_widths=[2.5, 3.0, 5.0, 6.5],
    row_colors=[COLOR_GREEN_ROW, COLOR_BLUE_ROW, COLOR_YELLOW_ROW, COLOR_RED_ROW],
)

add_h2(doc, "7.2  PHQ-A  (0–27 คะแนน)  —  อายุ 11–20 ปี")
make_table(doc,
    ["คะแนน", "ระดับ", "ป้าย", "Suicide Risk"],
    [
        ["0–4",   "none",        "ไม่มีภาวะซึมเศร้า", "ขึ้นกับ Q9/BQ1/BQ2"],
        ["5–9",   "mild",        "เล็กน้อย",           "ขึ้นกับ Q9/BQ1/BQ2"],
        ["10–14", "moderate",    "ปานกลาง",             "ขึ้นกับ Q9/BQ1/BQ2"],
        ["15–19", "severe",      "มาก",                 "ขึ้นกับ Q9/BQ1/BQ2"],
        ["20–27", "very_severe", "รุนแรง",              "ขึ้นกับ Q9/BQ1/BQ2"],
    ],
    col_widths=[2.5, 3.0, 5.0, 6.5],
    row_colors=[COLOR_GREEN_ROW, COLOR_BLUE_ROW, COLOR_YELLOW_ROW, COLOR_RED_ROW, COLOR_RED_ROW],
)

add_h2(doc, "7.3  CDI  (0–54 คะแนน)  —  อายุต่ำกว่า 11 ปี")
make_table(doc,
    ["คะแนน", "ระดับ", "ป้าย", "Suicide Risk"],
    [
        ["0–14", "normal",   "ปกติ",               "ขึ้นกับ Q9 เท่านั้น"],
        ["≥ 15", "clinical", "ควรได้รับการดูแล",   "ขึ้นกับ Q9 เท่านั้น"],
    ],
    col_widths=[2.5, 3.0, 5.0, 6.5],
    row_colors=[COLOR_GREEN_ROW, COLOR_RED_ROW],
)

add_divider(doc)

# ==========================================
#  ส่วนที่ 8 — ประเด็นขอความเห็น
# ==========================================
add_h1(doc, "ประเด็นขอความเห็นจากจิตแพทย์/นักจิตวิทยา", "8")

add_note_box(doc,
    "กรุณาระบุความเห็นในแต่ละข้อด้านล่าง หรือเขียนข้อเสนอแนะเพิ่มเติมในช่องว่างที่จัดไว้ให้",
    COLOR_INFO_BOX, "📋")

add_h2(doc, "8.1  ST-5")
add_bullet(doc, "เกณฑ์คะแนน 0–4 / 5–7 / 8–11 / 12–15 สอดคล้องกับฉบับมาตรฐานกรมสุขภาพจิตหรือไม่?")
add_bullet(doc, "ข้อคำถาม 5 ข้อครอบคลุมมิติความเครียดที่สำคัญสำหรับนักเรียนเพียงพอหรือไม่?")
add_body(doc, "ความเห็น: .............................................................................................................................", indent=0.5)
add_body(doc, "....................................................................................................................................................", indent=0.5)

add_h2(doc, "8.2  PHQ-A")
add_bullet(doc, "เกณฑ์คะแนน (0–4 / 5–9 / 10–14 / 15–19 / 20–27) ตรงกับมาตรฐาน PHQ-A ฉบับสถาบันสุขภาพจิตเด็กฯ พ.ศ. 2561 หรือไม่?")
add_bullet(doc, "การตั้งค่า Suicide Risk ที่ Q9 ≥ 1 (แม้เพียง \"มีบางวัน\") เหมาะสมหรือควรปรับเป็น ≥ 2 หรือ ≥ 3?")
add_bullet(doc, "BQ1 และ BQ2 ที่ trigger suicide_risk โดยไม่นับคะแนน — เหมาะสมหรือไม่?")
add_bullet(doc, "การแสดง BQ1/BQ2 ต่อนักเรียนที่คะแนนต่ำมาก (none) มีความเสี่ยงทางคลินิกหรือไม่?")
add_body(doc, "ความเห็น: .............................................................................................................................", indent=0.5)
add_body(doc, "....................................................................................................................................................", indent=0.5)

add_h2(doc, "8.3  CDI")
add_bullet(doc, "เกณฑ์ cut-off ที่คะแนน ≥ 15 = clinical ตรงกับมาตรฐาน CDI ฉบับแปลไทยหรือไม่?")
add_bullet(doc, "การแบ่ง Group A / Group B (การกลับทิศทางคะแนน) ถูกต้องตามต้นฉบับหรือไม่? (โปรดตรวจสอบตารางข้อที่ 5.3)")
add_bullet(doc, "CDI Q9 ตอบ \"ข\" (คิดเรื่องตายแต่ไม่อยากทำ) ไม่ถือเป็น Suicide Risk — เห็นด้วยหรือไม่?")
add_bullet(doc, "CDI มีแค่ 2 ระดับ (normal/clinical) — ควรเพิ่มระดับกลางหรือไม่?")
add_body(doc, "ความเห็น: .............................................................................................................................", indent=0.5)
add_body(doc, "....................................................................................................................................................", indent=0.5)

add_h2(doc, "8.4  ระบบโดยรวม")
add_bullet(doc, "กฎการแบ่งอายุ (7–17 ปี → CDI / ≥ 18 ปี → PHQ-A) สอดคล้องกับบริบทโรงเรียนในพื้นที่หรือไม่?")
add_bullet(doc, "คำแนะนำเบื้องต้นที่แสดงต่อนักเรียน (ส่วนที่ 3.4, 4.5, 5.4) เหมาะสมทางคลินิกหรือไม่?")
add_bullet(doc, "สายด่วน 1323 และ 1387 เพียงพอหรือควรเพิ่มช่องทางอื่น?")
add_body(doc, "ความเห็น: .............................................................................................................................", indent=0.5)
add_body(doc, "....................................................................................................................................................", indent=0.5)

add_divider(doc)

# ==========================================
#  ภาคผนวก
# ==========================================
add_h1(doc, "ภาคผนวก — ตัวอย่างการคำนวณคะแนน")

add_h2(doc, "ตัวอย่าง ST-5")
make_table(doc,
    ["ข้อ", "Q1", "Q2", "Q3", "Q4", "Q5", "รวม", "ระดับ"],
    [["คะแนน", "2", "1", "3", "1", "2", "9", "เครียดปานกลาง (moderate)"]],
    col_widths=[2.5, 1.5, 1.5, 1.5, 1.5, 1.5, 1.5, 7.0],
    row_colors=[COLOR_YELLOW_ROW],
)

add_h2(doc, "ตัวอย่าง PHQ-A")
make_table(doc,
    ["ข้อ", "Q1", "Q2", "Q3", "Q4", "Q5", "Q6", "Q7", "Q8", "Q9", "BQ1", "BQ2", "รวม", "ระดับ", "Suicide"],
    [["คะแนน", "2", "1", "2", "1", "0", "1", "1", "0", "1", "ไม่มี", "ไม่เคย", "9", "mild", "ใช่ (Q9=1)"]],
    col_widths=[2.0, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8, 1.3, 1.5, 1.0, 1.5, 2.0],
    row_colors=[COLOR_YELLOW_ROW],
    font_size=9,
)
add_note_box(doc, "แม้ระดับจะเป็น mild (คะแนนรวม 9) แต่ Q9 = 1 ทำให้ Suicide Risk = ใช่ → ระบบแจ้งเตือนทันที", COLOR_WARN_BOX, "⚠️")

add_h2(doc, "ตัวอย่าง CDI")
make_table(doc,
    ["ข้อ", "กลุ่ม", "คำตอบ", "คะแนน", "หมายเหตุ"],
    [
        ["Q1",  "A", "ข",  "1", "ก=0, ข=1, ค=2"],
        ["Q2",  "B", "ข",  "1", "ก=2, ข=1, ค=0 (กลับทิศ)"],
        ["Q7",  "B", "ก",  "2", "ก=2 เพราะเป็นกลุ่ม B"],
        ["Q9",  "A", "ก",  "0", "ตอบ ก = ไม่คิดเรื่องตาย → Suicide Risk = ไม่มี"],
        ["...", "–", "–", "–",  "คะแนนรวมสมมุติ = 16"],
    ],
    col_widths=[1.5, 2.0, 2.0, 2.0, 9.5],
    row_colors=[None, COLOR_BLUE_ROW, COLOR_YELLOW_ROW, COLOR_GREEN_ROW, None],
)
add_body(doc, "คะแนนรวม 16 → ระดับ clinical (≥15)  |  Suicide Risk = ไม่มี (Q9 ตอบ ก)")

doc.add_paragraph()
add_divider(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("เอกสารนี้จัดทำเพื่อการตรวจสอบทางคลินิกโดยเฉพาะ — ระบบ LEMCS จังหวัดเลย")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r.font.name = "TH Sarabun New"

output_path = r"D:\@LEMCS\docs\LEMCS_Clinical_Review_Manual.docx"
doc.save(output_path)
print(f"OK: {output_path}")
